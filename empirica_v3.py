# ============================================================================
# EMPIRICA v2.0.0 - Complete Research Pipeline
# ============================================================================
# v1.0.0: MVP — World Bank, Semantic Scholar, PubMed, 7 agents, Streamlit UI
# v1.1.0: Model upgrade (Sonnet 4.5), extended thinking, dual literature queries,
#         academic paper formatting (margins, spacing, page numbers, title page)
# v1.2.0: MECE policy section (bold lead sentences), conclusion/policy split
# v1.3.0: OMML equations in Word (native equation objects), UI overhaul
# v1.4.0: AMECO dataset via DBnomics — EU macro/fiscal hypotheses auto-route
# v1.5.0: Academic writing discipline overhaul — paragraph architecture, section
#          split (methodology/results/discussion), cartographic lit review,
#          anti-editorializing rules, causality language enforcement
# v1.6.0: Orwell + McCloskey + academic populist rewrite — dead metaphor ban,
#          cut filler words, short words over long, active voice, stress-at-end,
#          no throat-clearing, coefficient translation for non-specialists,
#          top-papers intro pattern (consensus→disruption→this paper→preview)
# v1.7.0: Section-specific structural fixes — intro opens with policy problem,
#          methodology justifies FE choice, results narrate across specs,
#          discussion adds "what data cannot answer" paragraph
# v1.8.0: Per-section temperatures (0.1→0.5), IMF intro/policy style,
#          PubMed structured abstract, finding-first lit review citations
# v1.9.0: Anti-AI cadence overhaul — few-shot voice examples from Hanushek &
#          Woessmann (2021) embedded in prompts, sentence-length variance rules,
# v2.0.0: IMF style consolidation - writing rules simplified to Orwell + IMF only,
#          few-shots from Shibata (2025 IMF) and Hanushek/Woessmann (2021),
#          abstract citations banned, descriptive stats table fixed (min/max + controls),
#          banned words list (drive/drives/driven etc), banned punctuation (em dash,
#          semicolons, colons for drama), hyphens only, indicator codes banned from prose,
#          proofreader overhauled with punctuation enforcement.
#          rough transition instructions, author-subordination enforcement,
#          proofreader AI-cadence detection (monotone rhythm, smooth transitions,
#          hedge stacking, evaluative filler), parenthetical aside encouragement
#
# Usage:
#   As module (from Streamlit/app.py):
#       from empirica_v3 import run_empirica
#       run_empirica("Your hypothesis here")
#
#   As standalone script:
#       export ANTHROPIC_API_KEY=sk-ant-your-key
#       python empirica_v3.py "Your hypothesis here"
# ============================================================================

import os
import sys
import json
import re
import time
import warnings
from datetime import datetime

import requests
import numpy as np
import pandas as pd
import scipy.stats as scipystats
import statsmodels.api as sm
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

import anthropic

try:
    from dbnomics import fetch_series as dbnomics_fetch
    HAS_DBNOMICS = True
except ImportError:
    HAS_DBNOMICS = False
    print("⚠️  dbnomics not installed — AMECO source unavailable. pip install dbnomics")

warnings.filterwarnings("ignore")


# ============================================================================
# CONFIGURATION
# ============================================================================
CLAUDE_MODEL = "claude-opus-4-6"                      # ← Opus 4.6
OUTPUT_DIR = "output"

INDICATOR_FAMILIES = {
    "NY.GDP": "GDP",
    "SE.XPD": "Education spending",
    "SH.XPD": "Health spending",
    "SP.DYN": "Demographics",
    "SI.POV": "Poverty",
}


# ============================================================================
# CLAUDE API HELPERS
# ============================================================================
def get_claude_client():
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY not set.")
    return anthropic.Anthropic(api_key=api_key)


def ask_claude(system: str, user: str, max_tokens: int = 4000, temperature: float = 0.3,
               extended_thinking: bool = False, thinking_budget: int = 10000) -> str:
    """Call Claude API. If extended_thinking=True, uses thinking with given budget."""
    client = get_claude_client()

    kwargs = dict(
        model=CLAUDE_MODEL,
        system=system,
        messages=[{"role": "user", "content": user}],
    )

    if extended_thinking:                              # ← CHANGE 2: extended thinking
        kwargs["temperature"] = 1                      #   thinking requires temperature=1
        kwargs["thinking"] = {"type": "enabled", "budget_tokens": thinking_budget}
        kwargs["max_tokens"] = max_tokens + thinking_budget
    else:
        kwargs["temperature"] = temperature
        kwargs["max_tokens"] = max_tokens

    response = client.messages.create(**kwargs)

    # Extract text (skip thinking blocks)
    for block in response.content:
        if block.type == "text":
            return block.text
    return ""


def ask_claude_json(system: str, user: str, max_tokens: int = 4000, temperature: float = 0.3,
                    extended_thinking: bool = False, thinking_budget: int = 10000) -> dict:
    raw = ask_claude(system, user, max_tokens, temperature, extended_thinking, thinking_budget)
    cleaned = re.sub(r"```json\s*", "", raw)
    cleaned = re.sub(r"```\s*", "", cleaned)
    cleaned = cleaned.strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if match:
            return json.loads(match.group())
        raise ValueError(f"Could not parse JSON from Claude response:\n{raw[:500]}")


def strip_markdown(text: str) -> str:
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"\$\$[\s\S]*?\$\$", "", text)
    text = re.sub(r"\$(.+?)\$", r"\1", text)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def strip_duplicate_heading(text: str, heading: str) -> str:
    lines = text.strip().split("\n")
    if not lines:
        return text
    first = lines[0].strip().lower()
    heading_clean = heading.strip().lower()
    first_no_num = re.sub(r"^\d+[\.\)]\s*", "", first)
    heading_no_num = re.sub(r"^\d+[\.\)]\s*", "", heading_clean)
    if first_no_num == heading_no_num or first == heading_no_num:
        return "\n".join(lines[1:]).strip()
    return text


# ============================================================================
# REGION MAPPING (for colored scatterplots)
# ============================================================================
def fetch_country_regions() -> dict:
    try:
        resp = requests.get(
            "https://api.worldbank.org/v2/country?format=json&per_page=300",
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        if len(data) < 2:
            return {}
        mapping = {}
        for c in data[1]:
            code = c.get("id", "")
            region = c.get("region", {}).get("value", "Other")
            name = c.get("name", "")
            if region and region != "Aggregates":
                mapping[name] = region
        return mapping
    except Exception:
        return {}


REGION_COLORS = {
    "East Asia & Pacific": "#E63946",
    "Europe & Central Asia": "#457B9D",
    "Latin America & Caribbean": "#2A9D8F",
    "Middle East & North Africa": "#E9C46A",
    "North America": "#264653",
    "South Asia": "#F4A261",
    "Sub-Saharan Africa": "#6A0572",
    "Other": "#AAAAAA",
}


# ============================================================================
# CHART GENERATION (ggplot-style)
# ============================================================================
def setup_ggplot_style():
    plt.rcParams.update({
        "figure.facecolor": "white",
        "axes.facecolor": "#F0F0F0",
        "axes.grid": True,
        "grid.color": "white",
        "grid.linewidth": 1.2,
        "axes.edgecolor": "#CCCCCC",
        "axes.linewidth": 0.8,
        "font.family": "sans-serif",
        "font.size": 10,
        "axes.titlesize": 13,
        "axes.labelsize": 11,
        "xtick.labelsize": 9,
        "ytick.labelsize": 9,
        "legend.fontsize": 8,
        "figure.dpi": 150,
    })


def generate_scatterplot(df: pd.DataFrame, plan: dict, output_dir: str) -> str:
    setup_ggplot_style()
    print("  📊 Generating scatterplot...")

    regions = fetch_country_regions()
    df_plot = df.copy()
    df_plot["region"] = df_plot["country"].map(regions).fillna("Other")

    country_means = df_plot.groupby(["country", "region"]).agg(
        x=("x", "mean"), y=("y", "mean")
    ).reset_index()

    fig, ax = plt.subplots(figsize=(8, 5.5))

    for region, color in REGION_COLORS.items():
        subset = country_means[country_means["region"] == region]
        if len(subset) > 0:
            ax.scatter(
                subset["x"], subset["y"],
                c=color, label=region, alpha=0.7, s=35, edgecolors="white", linewidths=0.4,
            )

    valid = country_means.dropna(subset=["x", "y"])
    if len(valid) > 2:
        z = np.polyfit(valid["x"], valid["y"], 1)
        p = np.poly1d(z)
        x_line = np.linspace(valid["x"].min(), valid["x"].max(), 100)
        ax.plot(x_line, p(x_line), color="#333333", linewidth=1.5, linestyle="--", alpha=0.7)

    ax.set_xlabel(plan["x_label"])
    ax.set_ylabel(plan["y_label"])
    ax.set_title(f"{plan['x_label']} vs {plan['y_label']}", fontweight="bold")
    ax.legend(loc="upper left", frameon=True, facecolor="white", edgecolor="#CCCCCC", ncol=2)

    plt.tight_layout()
    path = os.path.join(output_dir, "scatterplot.png")
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"    ✅ Saved: {path}")
    return path


def generate_coefficient_plot(results: dict, plan: dict, output_dir: str) -> str:
    setup_ggplot_style()
    print("  📊 Generating coefficient plot...")

    specs = []
    labels = []

    if "ols_controls" in results and "error" not in results["ols_controls"]:
        r = results["ols_controls"]
        specs.append((r["coefficient"], r["std_error"], r["p_value"]))
        labels.append("OLS + Controls")

    if "fixed_effects" in results and "error" not in results["fixed_effects"]:
        r = results["fixed_effects"]
        specs.append((r["coefficient"], r["std_error"], r["p_value"]))
        labels.append("Fixed Effects")

    if not specs:
        if "ols" in results and "error" not in results["ols"]:
            r = results["ols"]
            specs.append((r["coefficient"], r["std_error"], r["p_value"]))
            labels.append("OLS")

    if not specs:
        return ""

    fig, ax = plt.subplots(figsize=(6, max(2.5, len(specs) * 1.2)))

    y_pos = range(len(specs))
    coefs = [s[0] for s in specs]
    errors = [s[1] * 1.96 for s in specs]
    colors = ["#2A9D8F" if s[2] < 0.05 else "#E76F51" for s in specs]

    ax.barh(y_pos, coefs, xerr=errors, color=colors, alpha=0.8, height=0.5,
            edgecolor="white", capsize=4)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels)
    ax.axvline(x=0, color="#333333", linewidth=0.8, linestyle="-")
    ax.set_xlabel(f"Effect on {plan['y_label']}")
    ax.set_title("Coefficient Estimates (95% CI)", fontweight="bold")

    from matplotlib.patches import Patch
    legend_elements = [
        Patch(facecolor="#2A9D8F", label="p < 0.05"),
        Patch(facecolor="#E76F51", label="p ≥ 0.05"),
    ]
    ax.legend(handles=legend_elements, loc="lower right", frameon=True, facecolor="white")

    plt.tight_layout()
    path = os.path.join(output_dir, "coefficients.png")
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"    ✅ Saved: {path}")
    return path


# ============================================================================
# TAUTOLOGY GUARD
# ============================================================================
def check_tautology(x_code: str, y_code: str) -> bool:
    if x_code == y_code:
        return True
    for prefix in INDICATOR_FAMILIES:
        if x_code.startswith(prefix) and y_code.startswith(prefix):
            return True
    x_parts = x_code.split(".")
    y_parts = y_code.split(".")
    if len(x_parts) >= 2 and len(y_parts) >= 2:
        if x_parts[0] == y_parts[0] and x_parts[1] == y_parts[1]:
            return True
    return False


# ============================================================================
# INDICATOR VALIDATION
# ============================================================================
def validate_indicator(indicator: str) -> dict:
    try:
        resp = requests.get(
            f"https://api.worldbank.org/v2/indicator/{indicator}?format=json",
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        if len(data) >= 2 and data[1]:
            info = data[1][0]
            return {
                "id": info.get("id", ""),
                "name": info.get("name", ""),
                "source": info.get("source", {}).get("value", ""),
            }
    except Exception:
        pass
    return None


def check_data_availability(indicator: str, start_year: int = 2000, end_year: int = 2023) -> int:
    try:
        resp = requests.get(
            f"https://api.worldbank.org/v2/country/all/indicator/{indicator}"
            f"?date={start_year}:{end_year}&format=json&per_page=1&page=1",
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        if len(data) >= 1 and isinstance(data[0], dict):
            return data[0].get("total", 0)
    except Exception:
        pass
    return 0


def search_wb_indicators(keyword: str, max_results: int = 5) -> list:
    try:
        resp = requests.get(
            f"https://api.worldbank.org/v2/indicator?format=json&per_page=100",
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        if len(data) < 2 or not data[1]:
            return []
        kw = keyword.lower()
        matches = []
        for ind in data[1]:
            name = ind.get("name", "").lower()
            if kw in name:
                matches.append({
                    "code": ind["id"],
                    "name": ind["name"],
                })
        return matches[:max_results]
    except Exception:
        return []


def validate_and_fix_indicators(plan: dict) -> dict:
    print("  🔍 Validating indicators...")

    for var_key, label_key in [("independent_var", "x_label"), ("dependent_var", "y_label")]:
        code = plan[var_key]
        info = validate_indicator(code)

        if not info:
            print(f"    ⚠️  {code} does not exist in World Bank!")
            alt = ask_claude_json(
                system="You are a World Bank data expert. Suggest a VALID World Bank indicator code. Return JSON: {\"code\": \"XX.XXX.XXX\", \"name\": \"description\"}",
                user=f"The indicator {code} ({plan[label_key]}) does not exist. Suggest a valid alternative that measures the same concept.",
            )
            plan[var_key] = alt.get("code", code)
            plan[label_key] = alt.get("name", plan[label_key])
            print(f"    ✅ Replaced with: {plan[var_key]} ({plan[label_key]})")
        else:
            count = check_data_availability(code, plan.get("start_year", 2000), plan.get("end_year", 2023))
            if count < 200:
                print(f"    ⚠️  {code} has very sparse data ({count} points). Asking AI for denser alternative...")
                alt = ask_claude_json(
                    system="""You are a World Bank data expert. The user needs an indicator with GOOD data coverage (most countries, most years).
Suggest a VALID World Bank indicator that measures the same concept but has better data availability.
Common well-populated indicators include:
- NY.GDP.PCAP.PP.KD, NY.GDP.MKTP.KD.ZG, SP.DYN.LE00.IN, SP.DYN.IMRT.IN
- SE.XPD.TOTL.GD.ZS, SH.XPD.CHEX.GD.ZS, IT.NET.USER.ZS, SP.URB.TOTL.IN.ZS
- SL.UEM.TOTL.ZS, FP.CPI.TOTL.ZG, SP.POP.GROW, EG.ELC.ACCS.ZS
Return JSON: {"code": "XX.XXX.XXX", "name": "description", "reasoning": "why this is better"}""",
                    user=f"Indicator {code} ({plan[label_key]}) has only {count} data points (very sparse). I need something that measures '{plan[label_key]}' but with much better coverage across countries and years.",
                )
                new_code = alt.get("code", code)
                new_count = check_data_availability(new_code, plan.get("start_year", 2000), plan.get("end_year", 2023))
                if new_count > count:
                    plan[var_key] = new_code
                    plan[label_key] = alt.get("name", plan[label_key])
                    print(f"    ✅ Switched to: {new_code} ({plan[label_key]}) — {new_count} data points")
                else:
                    print(f"    ℹ️  Keeping {code} — alternative wasn't better")
            else:
                print(f"    ✅ {code} — {count} data points (good)")

    return plan


# ============================================================================
# AGENT 1: HYPOTHESIS PARSER (AI - extended thinking)
# ============================================================================
def ai_parse_hypothesis(hypothesis_text: str) -> dict:
    print("\n🧠 AGENT 1: Parsing hypothesis with AI (extended thinking)...")

    plan = ask_claude_json(
        system="""You are a research methodology expert with deep knowledge of the World Bank's data catalog (16,000+ indicators) AND the European Commission's AMECO database (481 datasets, 40+ EU/OECD countries).

Given a hypothesis, decide the BEST data source and pick indicator codes.

DATA SOURCE SELECTION:
- Use "worldbank" for: global/developing country topics, health, education, poverty, environment, infrastructure, demographics
- Use "ameco" for: EU/euro area macro-fiscal topics, fiscal policy, output gaps, structural deficits, unit labour costs, government debt, inflation (HICP), unemployment, current account, potential GDP, cyclical adjustment
- Use "both" only if the hypothesis explicitly compares EU vs global data (rare)

AMECO DATASET CODES (via DBnomics, provider="AMECO"):
  ZUTN — Unemployment rate (total)
  UVGD — GDP growth rate (real)
  OVGD — GDP at current prices
  AVGDGP — GDP per capita
  UBCA — Current account balance (% GDP)
  UICP — Inflation (HICP)
  UYCP — Inflation (GDP deflator)
  UDGG — Government gross debt (% GDP)
  UBLGE — Government balance (% GDP)
  UBLGAP — Structural budget balance (% potential GDP)
  PLCDQ — Nominal unit labour cost
  AMGN — Imports of goods and services
  AXGN — Exports of goods and services
  UOGG — Output gap (% potential GDP)
  OKND — Gross fixed capital formation (% GDP)
  NETD — Net lending/borrowing
  URTL — Long-term interest rate
  USTN — Short-term interest rate
AMECO dimensions: use {"geo": ["ea20"]} for euro area aggregate, or omit for all countries.
AMECO years: typically 1960-2025 (includes Commission forecasts).

WORLD BANK INDICATOR CODES (you know thousands, use any valid one):
WELL-POPULATED indicators (prefer these when possible):
GDP: NY.GDP.PCAP.PP.KD, NY.GDP.MKTP.KD.ZG, NY.GDP.PCAP.KD.ZG
Trade: NE.EXP.GNFS.ZS, NE.IMP.GNFS.ZS, TG.VAL.TOTL.GD.ZS
Finance: FP.CPI.TOTL.ZG, FM.LBL.BMNY.GD.ZS, BX.KLT.DINV.WD.GD.ZS
Education: SE.XPD.TOTL.GD.ZS, SE.SEC.ENRR, SE.TER.ENRR, SE.PRM.ENRR
Health: SH.XPD.CHEX.GD.ZS, SP.DYN.LE00.IN, SP.DYN.IMRT.IN, SH.MED.PHYS.ZS
Infrastructure: IT.NET.USER.ZS, EG.ELC.ACCS.ZS, IT.CEL.SETS.P2
Demographics: SP.URB.TOTL.IN.ZS, SP.POP.GROW, SP.DYN.TFRT.IN
Labor: SL.UEM.TOTL.ZS, SL.TLF.CACT.ZS, SL.AGR.EMPL.ZS
Governance: GE.EST, CC.EST, RL.EST, VA.EST
Environment: EN.ATM.CO2E.PC, EG.USE.ELEC.KH.PC, AG.LND.FRST.ZS
Poverty: SI.POV.DDAY (note: SI.POV.GINI has VERY sparse data, avoid it)
Water/Sanitation: SH.H2O.SMDW.ZS, SH.STA.SMSS.ZS

CRITICAL RULES:
1. X and Y MUST be from DIFFERENT domains, never two GDP indicators, two health indicators, etc.
2. The relationship must be CAUSAL/INTERESTING, not an accounting identity
3. PREFER indicators with GOOD data coverage — most countries, most years
4. Pick 2-4 control variables that are CONFOUNDERS (from the SAME source as X/Y)
5. For AMECO: control variables should also be AMECO dataset codes

Return JSON:
{
    "data_source": "worldbank" | "ameco" | "both",
    "title": "Academic paper title (specific, not generic)",
    "statement": "Cleaned hypothesis",
    "independent_var": "World Bank indicator code for X (ONLY if data_source is worldbank/both)",
    "dependent_var": "World Bank indicator code for Y (ONLY if data_source is worldbank/both)",
    "ameco_independent": {
        "dataset": "AMECO dataset code (e.g. UDGG)",
        "dimensions": {"geo": ["ea20"]},
        "label": "Human-readable label"
    },
    "ameco_dependent": {
        "dataset": "AMECO dataset code",
        "dimensions": {"geo": ["ea20"]},
        "label": "Human-readable label"
    },
    "x_label": "Human-readable label for X",
    "y_label": "Human-readable label for Y",
    "control_vars": [
        {"code": "indicator code OR AMECO dataset", "label": "label", "rationale": "why", "source": "worldbank|ameco"}
    ],
    "start_year": 2000,
    "end_year": 2023,
    "pubmed_query": "search query for PubMed (focused)",
    "pubmed_query_broad": "broader/different-angle PubMed query",
    "semantic_scholar_query": "search query for Semantic Scholar (focused)",
    "semantic_scholar_query_broad": "broader/different-angle Semantic Scholar query",
    "jel_codes": "JEL classification codes (e.g., O11, C23, I15)",
    "keywords": "4-6 keywords for the paper",
    "reasoning": "why these indicators and this data source are the best choice"
}

NOTE: If data_source is "worldbank", the ameco_* fields can be null/omitted.
If data_source is "ameco", independent_var/dependent_var can be null/omitted.""",
        user=f'Hypothesis: "{hypothesis_text}"\n\nPick the BEST indicators. Prefer well-populated ones. X = CAUSE, Y = EFFECT.\nAlso generate TWO search queries per database — one focused, one broader — to maximize literature coverage.',
        extended_thinking=True,           # ← CHANGE 2: thinking ON for Agent 1
        thinking_budget=10000,
    )

    source = plan.get("data_source", "worldbank")

    if source in ("worldbank", "both"):
        # Tautology check (World Bank codes only)
        if check_tautology(plan.get("independent_var", ""), plan.get("dependent_var", "")):
            print(f"  ⚠️  TAUTOLOGY DETECTED: {plan['independent_var']} -> {plan['dependent_var']}")

            h = hypothesis_text.lower()
            if "health" in h and ("life expectancy" in h or "mortality" in h or "life" in h):
                plan["independent_var"] = "SH.XPD.CHEX.GD.ZS"
                plan["dependent_var"] = "SP.DYN.LE00.IN"
                plan["x_label"] = "Current health expenditure (% of GDP)"
                plan["y_label"] = "Life expectancy at birth (years)"
            elif "education" in h and ("gdp" in h or "growth" in h or "income" in h):
                plan["independent_var"] = "SE.XPD.TOTL.GD.ZS"
                plan["dependent_var"] = "NY.GDP.PCAP.PP.KD"
                plan["x_label"] = "Government expenditure on education (% of GDP)"
                plan["y_label"] = "GDP per capita (PPP, constant 2017 $)"
            elif "internet" in h and ("gdp" in h or "growth" in h or "income" in h):
                plan["independent_var"] = "IT.NET.USER.ZS"
                plan["dependent_var"] = "NY.GDP.PCAP.PP.KD"
                plan["x_label"] = "Individuals using the Internet (% of population)"
                plan["y_label"] = "GDP per capita (PPP, constant 2017 $)"

            print(f"  ✅ Corrected to: {plan['x_label']} -> {plan['y_label']}")

        # Default controls fallback (World Bank)
        if len(plan.get("control_vars", [])) < 2:
            default_controls = [
                {"code": "NY.GDP.PCAP.PP.KD", "label": "GDP per capita (PPP)", "rationale": "Income level confounder", "source": "worldbank"},
                {"code": "SE.SEC.ENRR", "label": "Secondary school enrollment", "rationale": "Education confounder", "source": "worldbank"},
                {"code": "SP.URB.TOTL.IN.ZS", "label": "Urban population (%)", "rationale": "Urbanization confounder", "source": "worldbank"},
            ]
            existing_codes = {c["code"] for c in plan.get("control_vars", [])}
            for dc in default_controls:
                if dc["code"] not in existing_codes and dc["code"] != plan.get("independent_var") and dc["code"] != plan.get("dependent_var"):
                    plan.setdefault("control_vars", []).append(dc)
                    if len(plan["control_vars"]) >= 3:
                        break

        x_code = plan.get('independent_var', '?')
        y_code = plan.get('dependent_var', '?')
    else:
        # AMECO source — labels from ameco_independent/dependent
        ameco_x = plan.get("ameco_independent", {})
        ameco_y = plan.get("ameco_dependent", {})
        plan.setdefault("x_label", ameco_x.get("label", ameco_x.get("dataset", "X")))
        plan.setdefault("y_label", ameco_y.get("label", ameco_y.get("dataset", "Y")))
        x_code = f"AMECO/{ameco_x.get('dataset', '?')}"
        y_code = f"AMECO/{ameco_y.get('dataset', '?')}"

    print(f"  -> Source: {source}")
    print(f"  -> Title: {plan['title']}")
    print(f"  -> X: {plan['x_label']} ({x_code})")
    print(f"  -> Y: {plan['y_label']} ({y_code})")
    print(f"  -> Controls: {', '.join(c['label'] for c in plan.get('control_vars', []))}")
    print(f"  -> Years: {plan['start_year']}-{plan['end_year']}")

    if source in ("worldbank", "both"):
        plan = validate_and_fix_indicators(plan)

    return plan


# ============================================================================
# AGENT 2: DATA COLLECTOR (Code)
# ============================================================================
class WorldBankFetcher:
    BASE_URL = "https://api.worldbank.org/v2"

    AGGREGATES = {
        "WLD", "HIC", "LIC", "LMC", "MIC", "UMC", "LMY", "HPC",
        "EAS", "ECS", "LCN", "MEA", "NAC", "SAS", "SSF", "AFE",
        "AFW", "ARB", "CEB", "CSS", "EAP", "EAR", "EMU", "EUU",
        "FCS", "IDA", "IDX", "LAC", "LDC", "LTE", "MNA", "OED",
        "OSS", "PRE", "PSS", "PST", "SSA", "SST", "TEA", "TEC",
        "TLA", "TMN", "TSA", "TSS", "IBD", "IBT", "IDB",
    }

    def fetch(self, indicator: str, start_year: int, end_year: int) -> pd.DataFrame:
        print(f"  📊 Fetching {indicator} ({start_year}-{end_year})...")
        all_data = []
        page = 1
        while True:
            url = (
                f"{self.BASE_URL}/country/all/indicator/{indicator}"
                f"?date={start_year}:{end_year}&format=json&per_page=1000&page={page}"
            )
            resp_data = None
            for attempt in range(3):
                try:
                    resp = requests.get(url, timeout=45)
                    resp.raise_for_status()
                    resp_data = resp.json()
                    break
                except Exception as e:
                    if attempt < 2:
                        print(f"    ⚠️  Retry {attempt + 1}/3 for {indicator}: {e}")
                        time.sleep(2)
                    else:
                        print(f"    ⚠️  World Bank API failed after 3 attempts: {e}")

            if not resp_data:
                break

            if len(resp_data) < 2 or not resp_data[1]:
                break

            for record in resp_data[1]:
                value = record.get("value")
                if value is not None:
                    cc = record.get("country", {}).get("id", "")
                    if cc not in self.AGGREGATES:
                        all_data.append({
                            "country": record["country"]["value"],
                            "country_code": cc,
                            "year": int(record["date"]),
                            "value": float(value),
                        })

            if page >= resp_data[0].get("pages", 1):
                break
            page += 1

        df = pd.DataFrame(all_data)
        if not df.empty:
            print(f"    ✅ {len(df)} observations, {df['country'].nunique()} countries")
        else:
            print(f"    ⚠️  No data returned for {indicator}")
        return df


class DBnomicsFetcher:
    """Fetches AMECO data via DBnomics. Returns same DataFrame shape as WorldBankFetcher."""
    PROVIDER = "AMECO"

    # AMECO geo codes → readable country names (common ones)
    GEO_LABELS = {
        "aut": "Austria", "bel": "Belgium", "bgr": "Bulgaria", "hrv": "Croatia",
        "cyp": "Cyprus", "cze": "Czechia", "dnk": "Denmark", "est": "Estonia",
        "fin": "Finland", "fra": "France", "deu": "Germany", "grc": "Greece",
        "hun": "Hungary", "irl": "Ireland", "ita": "Italy", "lva": "Latvia",
        "ltu": "Lithuania", "lux": "Luxembourg", "mlt": "Malta", "nld": "Netherlands",
        "pol": "Poland", "prt": "Portugal", "rou": "Romania", "svk": "Slovakia",
        "svn": "Slovenia", "esp": "Spain", "swe": "Sweden", "gbr": "United Kingdom",
        "usa": "United States", "jpn": "Japan", "can": "Canada", "aus": "Australia",
        "nor": "Norway", "che": "Switzerland", "isl": "Iceland", "kor": "South Korea",
        "nzl": "New Zealand", "mex": "Mexico", "tur": "Turkey",
        "ea20": "Euro Area (20)", "ea19": "Euro Area (19)", "eu27": "EU-27",
    }

    # Aggregates to exclude from country-level analysis (like WB aggregates)
    AGGREGATES = {"ea20", "ea19", "ea18", "eu27", "eu28", "eu15", "g7", "g20", "oecd"}

    def fetch(self, dataset_code: str, dimensions: dict = None,
              start_year: int = 2000, end_year: int = 2023,
              include_aggregates: bool = False) -> pd.DataFrame:
        """Fetch AMECO dataset via DBnomics. Returns DataFrame[country, country_code, year, value]."""
        if not HAS_DBNOMICS:
            print(f"    ⚠️  dbnomics not installed — cannot fetch AMECO/{dataset_code}")
            return pd.DataFrame()

        print(f"  📊 Fetching AMECO/{dataset_code} via DBnomics...")
        try:
            df = dbnomics_fetch(
                provider_code=self.PROVIDER,
                dataset_code=dataset_code,
                dimensions=dimensions or {}
            )

            if df.empty:
                print(f"    ⚠️  No data returned for AMECO/{dataset_code}")
                return pd.DataFrame()

            # Standardize columns to match WorldBankFetcher output
            df = df.rename(columns={"original_period": "year"})
            df["year"] = pd.to_numeric(df["year"], errors="coerce")
            df["value"] = pd.to_numeric(df["value"], errors="coerce")
            df = df.dropna(subset=["year", "value"])
            df["year"] = df["year"].astype(int)
            df = df[(df["year"] >= start_year) & (df["year"] <= end_year)]

            # Extract country from geo dimension
            geo_col = None
            for candidate in ["geo", "GEO", "country", "unit"]:
                if candidate in df.columns:
                    geo_col = candidate
                    break

            if geo_col is None:
                # Try to extract from series_code (format: provider/dataset/series)
                if "series_code" in df.columns:
                    df["country_code"] = df["series_code"].str.split(".").str[-1].str.lower()
                else:
                    print(f"    ⚠️  Cannot identify country column in AMECO/{dataset_code}")
                    return pd.DataFrame()
            else:
                df["country_code"] = df[geo_col].astype(str).str.lower()

            # Filter out aggregates unless requested
            if not include_aggregates:
                df = df[~df["country_code"].isin(self.AGGREGATES)]

            # Map codes to readable names
            df["country"] = df["country_code"].map(self.GEO_LABELS).fillna(df["country_code"].str.upper())

            df = df[["country", "country_code", "year", "value"]].reset_index(drop=True)
            print(f"    ✅ {len(df)} obs, {df['country'].nunique()} countries")
            return df

        except Exception as e:
            print(f"    ⚠️  DBnomics error for AMECO/{dataset_code}: {e}")
            return pd.DataFrame()


class SemanticScholarSearcher:
    BASE_URL = "https://api.semanticscholar.org/graph/v1"

    def search(self, query: str, max_results: int = 15) -> list:   # ← CHANGE 3: 15 default
        papers = []
        for attempt in range(3):
            try:
                print(f"  📖 Semantic Scholar (attempt {attempt + 1}): {query}")
                resp = requests.get(
                    f"{self.BASE_URL}/paper/search",
                    params={
                        "query": query, "limit": max_results,
                        "fields": "title,authors,year,journal,externalIds,abstract,citationCount",
                    },
                    timeout=30,
                )
                resp.raise_for_status()
                papers = resp.json().get("data", [])
                if papers:
                    break
            except Exception as e:
                print(f"    -> Attempt {attempt + 1} failed: {e}")
                if attempt < 2:
                    time.sleep(2)
                papers = []

        articles = []
        for p in papers:
            try:
                title = p.get("title", "")
                authors_raw = p.get("authors", [])
                authors = [a.get("name", "") for a in authors_raw if a.get("name")]
                year = str(p.get("year", ""))
                if not (title and authors and year and year != "None"):
                    continue
                journal_info = p.get("journal")
                journal = journal_info.get("name", "Unknown") if journal_info else "Unknown"
                ext_ids = p.get("externalIds", {}) or {}
                doi = ext_ids.get("DOI", "")
                abstract = (p.get("abstract") or "")[:500]
                citations = p.get("citationCount", 0) or 0
                authors_short = f"{authors[0]} et al." if len(authors) > 3 else ", ".join(authors)
                articles.append({
                    "title": title, "authors": authors, "authors_short": authors_short,
                    "year": year, "journal": journal, "doi": doi, "pmid": "",
                    "abstract": abstract, "citations": citations, "source": "Semantic Scholar",
                })
                print(f"    -> {authors_short} ({year}) [{citations} cites] - {title[:60]}...")
            except Exception:
                continue
        articles.sort(key=lambda a: a.get("citations", 0), reverse=True)
        return articles


class PubMedSearcher:
    BASE_URL = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

    def search(self, query: str, max_results: int = 10) -> list:    # ← CHANGE 3: 10 default
        print(f"  📖 PubMed search: {query}")
        try:
            search_resp = requests.get(
                f"{self.BASE_URL}/esearch.fcgi",
                params={"db": "pubmed", "term": query, "retmax": max_results, "retmode": "json", "sort": "relevance"},
                timeout=15,
            )
            search_resp.raise_for_status()
            ids = search_resp.json().get("esearchresult", {}).get("idlist", [])
            if not ids:
                print("    ⚠️  No PubMed results")
                return []

            fetch_resp = requests.get(
                f"{self.BASE_URL}/esummary.fcgi",
                params={"db": "pubmed", "id": ",".join(ids), "retmode": "json"},
                timeout=15,
            )
            fetch_resp.raise_for_status()
            results = fetch_resp.json().get("result", {})

            articles = []
            for pmid in ids:
                info = results.get(pmid, {})
                if not info or pmid == "uids":
                    continue
                title = info.get("title", "").rstrip(".")
                authors_raw = info.get("authors", [])
                authors = [a.get("name", "") for a in authors_raw if a.get("name")]
                if not authors:
                    continue
                year = info.get("pubdate", "")[:4]
                journal = info.get("source", "Unknown")
                doi_list = [x.get("value", "") for x in info.get("articleids", []) if x.get("idtype") == "doi"]
                doi = doi_list[0] if doi_list else ""
                authors_short = f"{authors[0]} et al." if len(authors) > 3 else ", ".join(authors)
                articles.append({
                    "title": title, "authors": authors, "authors_short": authors_short,
                    "year": year, "journal": journal, "doi": doi, "pmid": pmid,
                    "abstract": "", "citations": 0, "source": "PubMed",
                })
                print(f"    -> {authors_short} ({year}) - {title[:60]}...")
            return articles
        except Exception as e:
            print(f"    ⚠️  PubMed error: {e}")
            return []


class LiteratureSearcher:
    """Runs dual queries on both Semantic Scholar and PubMed for 20-30+ papers."""
    def __init__(self):
        self.ss = SemanticScholarSearcher()
        self.pm = PubMedSearcher()

    def search(self, plan: dict) -> list:                          # ← CHANGE 3: dual queries
        print("\n📚 AGENT 2b: Searching literature (dual queries)...")

        all_articles = []

        # --- Semantic Scholar: focused query (15 papers) ---
        ss_query = plan.get("semantic_scholar_query", plan["statement"])
        all_articles.extend(self.ss.search(ss_query, max_results=15))
        time.sleep(1)

        # --- Semantic Scholar: broad query (10 papers, different angle) ---
        ss_broad = plan.get("semantic_scholar_query_broad", "")
        if ss_broad and ss_broad != ss_query:
            print("  📖 Semantic Scholar (broad query)...")
            all_articles.extend(self.ss.search(ss_broad, max_results=10))
            time.sleep(1)

        # --- PubMed: focused query (10 papers) ---
        pm_query = plan.get("pubmed_query", plan["statement"])
        all_articles.extend(self.pm.search(pm_query, max_results=10))
        time.sleep(1)

        # --- PubMed: broad query (5 papers) ---
        pm_broad = plan.get("pubmed_query_broad", "")
        if pm_broad and pm_broad != pm_query:
            print("  📖 PubMed (broad query)...")
            all_articles.extend(self.pm.search(pm_broad, max_results=5))

        # Deduplicate by DOI and title
        seen_dois = set()
        seen_titles = set()
        combined = []
        for article in all_articles:
            doi = article.get("doi", "")
            title_lower = article.get("title", "").lower().strip()
            if doi and doi in seen_dois:
                continue
            if title_lower and title_lower in seen_titles:
                continue
            if doi:
                seen_dois.add(doi)
            if title_lower:
                seen_titles.add(title_lower)
            combined.append(article)

        combined.sort(key=lambda a: a.get("citations", 0), reverse=True)
        print(f"  ✅ {len(combined)} unique articles found (target: 20+)")
        return combined


# ============================================================================
# AGENT 3: DATA REVIEWER (AI - extended thinking)
# ============================================================================
def ai_review_data(df: pd.DataFrame, plan: dict) -> dict:
    print("\n🔍 AGENT 3: AI reviewing data quality (extended thinking)...")

    summary = {
        "rows": len(df),
        "countries": int(df["country"].nunique()) if "country" in df.columns else 0,
        "years": f"{int(df['year'].min())}-{int(df['year'].max())}" if "year" in df.columns else "N/A",
    }
    for col in ["x", "y"]:
        if col in df.columns:
            summary[f"{col}_stats"] = {
                "mean": round(float(df[col].mean()), 4),
                "std": round(float(df[col].std()), 4),
                "min": round(float(df[col].min()), 4),
                "max": round(float(df[col].max()), 4),
                "missing_pct": round(float(df[col].isna().mean() * 100), 2),
                "zeros_pct": round(float((df[col] == 0).mean() * 100), 2),
            }

    review = ask_claude_json(
        system="""You are a data quality analyst. Review this dataset and recommend cleaning.
Return JSON:
{
    "assessment": "brief quality assessment",
    "winsorize": true/false,
    "winsorize_percentile": 1 or 5,
    "exclude_zeros_x": true/false,
    "exclude_zeros_y": true/false,
    "min_observations_per_country": 3 or 5,
    "exclude_countries": [],
    "warnings": []
}""",
        user=f"Hypothesis: {plan['statement']}\n\n{json.dumps(summary, indent=2)}",
        extended_thinking=True,           # ← CHANGE 2: thinking ON for Agent 3
        thinking_budget=8000,
    )

    print(f"  -> Assessment: {review.get('assessment', 'N/A')}")
    for w in review.get("warnings", []):
        print(f"  ⚠️  {w}")
    return review


def apply_cleaning(df: pd.DataFrame, review: dict) -> pd.DataFrame:
    print("  🧹 Applying cleaning...")
    original_len = len(df)

    if review.get("winsorize", False):
        pct = review.get("winsorize_percentile", 1) / 100
        for col in ["x", "y"]:
            if col in df.columns:
                df[col] = df[col].clip(df[col].quantile(pct), df[col].quantile(1 - pct))

    if review.get("exclude_zeros_x") and "x" in df.columns:
        df = df[df["x"] != 0]
    if review.get("exclude_zeros_y") and "y" in df.columns:
        df = df[df["y"] != 0]

    min_obs = review.get("min_observations_per_country", 3)
    if "country" in df.columns:
        counts = df.groupby("country").size()
        df = df[df["country"].isin(counts[counts >= min_obs].index)]

    for c in review.get("exclude_countries", []):
        if "country" in df.columns:
            df = df[df["country"] != c]

    print(f"  ✅ {original_len} -> {len(df)} rows")
    return df


# ============================================================================
# AGENT 4: STATISTICS ENGINE (Code)
# ============================================================================
class StatisticsEngine:
    def run_all(self, df: pd.DataFrame, plan: dict) -> dict:
        print("\n📈 AGENT 4: Running statistical analysis...")
        results = {}

        for col in ["x", "y"]:
            df[col] = pd.to_numeric(df[col], errors="coerce")
        df = df.dropna(subset=["x", "y"])

        if len(df) < 10:
            return {"error": "Insufficient data (< 10 observations)"}

        results["ols"] = self._ols(df)

        control_cols = [c for c in df.columns if c.startswith("control_")]
        if control_cols:
            results["ols_controls"] = self._ols_controls(df, control_cols)

        if df["country"].nunique() > 5:
            results["fixed_effects"] = self._fixed_effects(df)

        results["correlation"] = self._correlation(df)

        results["descriptive"] = {
            "n_obs": len(df),
            "n_countries": int(df["country"].nunique()),
            "year_range": f"{int(df['year'].min())}-{int(df['year'].max())}",
            "x_mean": round(float(df["x"].mean()), 4),
            "x_std": round(float(df["x"].std()), 4),
            "y_mean": round(float(df["y"].mean()), 4),
            "y_std": round(float(df["y"].std()), 4),
            "x_stats": {
                "mean": round(float(df["x"].mean()), 3),
                "std": round(float(df["x"].std()), 3),
                "min": round(float(df["x"].min()), 3),
                "max": round(float(df["x"].max()), 3),
            },
            "y_stats": {
                "mean": round(float(df["y"].mean()), 3),
                "std": round(float(df["y"].std()), 3),
                "min": round(float(df["y"].min()), 3),
                "max": round(float(df["y"].max()), 3),
            },
        }
        # Add control variable stats
        control_cols = [c for c in df.columns if c.startswith("control_")]
        for cc in control_cols:
            cname = cc.replace("control_", "")
            results["descriptive"][f"ctrl_{cname}_stats"] = {
                "mean": round(float(df[cc].mean()), 3),
                "std": round(float(df[cc].std()), 3),
                "min": round(float(df[cc].min()), 3),
                "max": round(float(df[cc].max()), 3),
            }
        return results

    def _ols(self, df):
        print("  📐 OLS regression...")
        try:
            X = sm.add_constant(df["x"])
            model = sm.OLS(df["y"], X).fit()
            r = {
                "coefficient": round(float(model.params.iloc[1]), 6),
                "intercept": round(float(model.params.iloc[0]), 6),
                "std_error": round(float(model.bse.iloc[1]), 6),
                "t_stat": round(float(model.tvalues.iloc[1]), 4),
                "p_value": round(float(model.pvalues.iloc[1]), 6),
                "r_squared": round(float(model.rsquared), 4),
                "adj_r_squared": round(float(model.rsquared_adj), 4),
                "n_obs": int(model.nobs),
                "f_stat": round(float(model.fvalue), 4),
                "significant": float(model.pvalues.iloc[1]) < 0.05,
            }
            sig = "***" if r["p_value"] < 0.001 else "**" if r["p_value"] < 0.01 else "*" if r["p_value"] < 0.05 else ""
            print(f"    -> B = {r['coefficient']} (p = {r['p_value']}) {sig}, R2 = {r['r_squared']}")
            return r
        except Exception as e:
            print(f"    ⚠️  OLS failed: {e}")
            return {"error": str(e)}

    def _ols_controls(self, df, control_cols):
        print(f"  📐 OLS with {len(control_cols)} controls...")
        try:
            df_clean = df.dropna(subset=["x", "y"] + control_cols)
            X = sm.add_constant(df_clean[["x"] + control_cols])
            model = sm.OLS(df_clean["y"], X).fit()
            r = {
                "coefficient": round(float(model.params["x"]), 6),
                "std_error": round(float(model.bse["x"]), 6),
                "p_value": round(float(model.pvalues["x"]), 6),
                "r_squared": round(float(model.rsquared), 4),
                "adj_r_squared": round(float(model.rsquared_adj), 4),
                "n_obs": int(model.nobs),
                "significant": float(model.pvalues["x"]) < 0.05,
                "controls_used": control_cols,
            }
            print(f"    -> B = {r['coefficient']} (p = {r['p_value']}), R2 = {r['r_squared']}")
            return r
        except Exception as e:
            print(f"    ⚠️  OLS+controls failed: {e}")
            return {"error": str(e)}

    def _fixed_effects(self, df):
        print("  📐 Country fixed effects...")
        try:
            df_fe = df.copy()
            for col in ["x", "y"]:
                df_fe[f"{col}_dm"] = df_fe[col] - df_fe.groupby("country")[col].transform("mean")
            X = sm.add_constant(df_fe["x_dm"])
            model = sm.OLS(df_fe["y_dm"], X).fit()
            r = {
                "coefficient": round(float(model.params.iloc[-1]), 6),
                "std_error": round(float(model.bse.iloc[-1]), 6),
                "p_value": round(float(model.pvalues.iloc[-1]), 6),
                "r_squared_within": round(float(model.rsquared), 4),
                "n_obs": int(model.nobs),
                "n_countries": int(df["country"].nunique()),
                "significant": float(model.pvalues.iloc[-1]) < 0.05,
            }
            print(f"    -> B(FE) = {r['coefficient']} (p = {r['p_value']}), R2w = {r['r_squared_within']}")
            return r
        except Exception as e:
            print(f"    ⚠️  FE failed: {e}")
            return {"error": str(e)}

    def _correlation(self, df):
        print("  📐 Correlation...")
        try:
            pr, pp = scipystats.pearsonr(df["x"], df["y"])
            sr, sp = scipystats.spearmanr(df["x"], df["y"])
            r = {
                "pearson_r": round(float(pr), 4), "pearson_p": round(float(pp), 6),
                "spearman_r": round(float(sr), 4), "spearman_p": round(float(sp), 6),
            }
            print(f"    -> Pearson r = {r['pearson_r']}, Spearman rho = {r['spearman_r']}")
            return r
        except Exception as e:
            print(f"    ⚠️  Correlation failed: {e}")
            return {"error": str(e)}


# ============================================================================
# AGENT 5: RESULTS INTERPRETER (AI - extended thinking)
# ============================================================================
def ai_interpret_results(results: dict, plan: dict, advocacy_angle: str = "", advocacy_temperature: int = 1) -> dict:
    print("\n⚖️ AGENT 5: AI interpreting results (extended thinking)...")

    advocacy_block = ""
    if advocacy_angle and advocacy_angle.strip():
        if advocacy_temperature <= 3:
            advocacy_block = f'\nNote: the author leans toward "{advocacy_angle.strip()}". Keep interpretation honest but when results are ambiguous, lean slightly toward this framing.'
        elif advocacy_temperature <= 6:
            advocacy_block = f'\nNote: the author\'s perspective is "{advocacy_angle.strip()}". When results support this view, emphasize it. When they don\'t, frame as "nuanced" rather than contradictory. Still be honest about effect sizes.'
        else:
            advocacy_block = f'\nNote: the author strongly advocates for "{advocacy_angle.strip()}". Frame the main_finding to support this perspective where possible. Find the strongest angle. Still report actual numbers honestly.'

    interpretation = ask_claude_json(
        system=f"""You are an econometrics expert. Given statistical results, provide an honest assessment.
If R2 is 0.04, say "very weak." Do NOT oversell.{advocacy_block}
Return JSON:
{{
    "strength": "strong / moderate / weak / very weak / none",
    "direction": "positive / negative / unclear",
    "confidence": "high / moderate / low",
    "main_finding": "One sentence summary",
    "caveats": ["limitations"],
    "recommended_tone": "confident / cautious / very cautious / skeptical",
    "additional_tests_suggested": []
}}""",
        user=f"Hypothesis: {plan['statement']}\n\nResults:\n{json.dumps(results, indent=2, default=str)}",
        extended_thinking=True,           # ← CHANGE 2: thinking ON for Agent 5
        thinking_budget=8000,
    )
    print(f"  -> {interpretation.get('strength', '?')} | {interpretation.get('recommended_tone', '?')}")
    print(f"  -> {interpretation.get('main_finding', 'N/A')}")
    return interpretation


# ============================================================================
# AGENT 6: PAPER WRITER (AI - academic discipline, NO extended thinking)
# ============================================================================
WRITING_RULES = """WRITING RULES (follow strictly):

═══ ORWELL'S RULES ═══
1. Never use a dead metaphor. No "sheds light on," "level playing field," "at the end of the day," "the landscape of." If you have seen the phrase in ten other papers, delete it and say what you actually mean.
2. Never use a long word where a short one will do. "Use" not "utilize." "Show" not "demonstrate." "About" not "approximately." "But" not "however" (when starting a sentence).
3. If it is possible to cut a word out, always cut it out. Delete "it is important to note that," "it should be noted that," "in order to," "the fact that," "a growing body of literature suggests."
4. Never use the passive where you can use the active. "We estimate" not "it was estimated." Keep passive only for method conventions: "Standard errors are clustered by country."
5. Never use jargon where plain English will do. Technical terms get a one-line parenthetical on first use.
6. Break any of these rules sooner than write anything barbarous.

═══ IMF WORKING PAPER STYLE ═══
7. Numbered paragraphs (1., 2., 3., etc.) in the introduction. Each paragraph opens with a bold declarative takeaway sentence. Attach a number to every claim.
8. Every coefficient must be translated into real-world units. "A one-percentage-point increase in X corresponds to 0.34 additional units of Y" is clear. Regression notation alone is not.
9. One hedge per claim, maximum. State the finding, qualify it once, move on.
10. Results are reported, not narrated. Open with a table reference. State the coefficient and significance. Translate once. Compare specifications.
11. The conclusion must be statable in one sentence a non-economist can understand.

═══ VOICE AND CADENCE ═══
12. VARY SENTENCE LENGTH. Mix 8-word sentences with 35-word sentences. After a long sentence, drop a short one: "The results are clear." Or: "That assumption is wrong."
13. BE BLUNT. If the coefficient is near zero, write: "The effect is negligible." If R-squared is 0.03, write: "The model explains almost nothing."
14. ROUGH TRANSITIONS. Do not connect every paragraph with "Furthermore" or "Additionally." Just start the next thought.
15. SUBORDINATE CITATIONS. Never make an author name the grammatical subject unless attributing a specific intellectual move. The finding leads, the citation follows in parentheses.
16. ALLOW IMPERFECT RHYTHM. Some paragraphs should be 2 sentences. Some should be 6. Do not make every paragraph claim-evidence-qualification.
17. FOCUS ON INTERPRETATION. When discussing results, talk about economic meaning, not indicator codes. "Education spending" not "SE.XPD.TOTL.GD.ZS." The reader cares about what the numbers mean, not which database column they came from.

═══ FEW-SHOT EXAMPLES (match this cadence) ═══

EXAMPLE A (IMF introduction, Shibata 2025):
"Spain's significant per capita income gap with highest-income euro area economies and the United States primarily reflects a wide productivity shortfall. In 2024, Spain's income per capita in PPP terms stood nearly 40 and 16 percent below that of the US and the other three largest euro area economies, respectively. While both lower capital intensity and fewer total working hours accounted for some of the gap, weaker total factor productivity accounted for over two-thirds of it."
Note: First sentence is the bold takeaway. Second sentence attaches numbers immediately. Third sentence qualifies with "while." No em dashes, no semicolons, no colons for drama.

EXAMPLE B (Hanushek and Woessmann 2021, introduction):
"Economic growth determines the well-being of society over time, but the source of differences in growth rates of countries are continuously debated. Intensive analysis over the past three decades has, however, produced a much clearer picture. In simplest terms, long-term economic growth is largely determined by the skills of the population."
Note: 3 sentences. The third is the punchline, short and declarative. The transition "In simplest terms" is conversational.

EXAMPLE C (IMF, discussing firm dynamics):
"The Spanish economy is less dynamic than European peers and, most strikingly, the US, particularly with respect to the footprint of young high-growth firms, resulting in an overabundance of small firms. While average entry and exit rates are broadly comparable, Spanish and European firms enter small. Even more importantly, Spanish and European firms struggle to scale up."
Note: "Even more importantly" is direct. No "Furthermore." The short sentence "Spanish and European firms enter small" punches after the long opener.

EXAMPLE D (Hanushek and Woessmann 2021, stating a null result):
"The relationship is now flat. School attainment is not statistically significant in the presence of the direct cognitive-skill measure of human capital."
Note: Blunt. No hedging. Two sentences where AI would write one long hedged one.

═══ BANNED WORDS AND PUNCTUATION ═══
- NEVER use: "drive," "drives," "driven," "driving" (say "is associated with," "corresponds to," "leads to" only with causal evidence). Also banned: "landscape," "paradigm," "robust" (as praise), "compelling," "striking," "remarkable," "notable," "critical" (as emphasis), "garnered," "pivotal," "key insight," "mounting evidence," "increasingly recognized," "underscores," "highlights."
- NEVER use em dashes. Use commas or parentheses instead.
- NEVER use semicolons for rhetorical effect. Use a period and start a new sentence.
- NEVER use colons for dramatic reveals.
- Use hyphens (-) for compound words. Not en dashes, not em dashes. Just hyphens.
- Never editorialize: no "the data speak clearly," "the gap tells a revealing story," "the punchline is cautionary."
- Never use rhetorical questions or address the reader directly.
- Never open with "Conventional wisdom holds..." or any hook/framing device.
- Never discuss channels or variables NOT measured in the dataset.
- No markdown formatting. No #, **, *, `, $$.
- Do NOT write a full paper. Write ONLY the section requested.

EQUATIONS:
- Wrap each equation in [EQ]...[/EQ] markers on its own line. Use Greek Unicode (α, β, γ, ε, μ, σ, δ, λ, θ, ρ, τ, φ) and _{subscript} for subscripts. Example: [EQ]Y_{it} = α + β × X_{it} + γ Z_{it} + ε_{it}[/EQ]"""


def build_advocacy_instruction(advocacy_angle: str, advocacy_temperature: int) -> str:
    """Generate prompt instructions for advocacy framing, scaled by temperature (1-10).
    Returns empty string if no advocacy angle is set."""
    if not advocacy_angle or not advocacy_angle.strip():
        return ""

    angle = advocacy_angle.strip()

    if advocacy_temperature <= 3:
        # Subtle - barely noticeable lean
        return f"""
FRAMING GUIDANCE (subtle):
The author has a perspective: "{angle}".
Incorporate this lightly. The paper should read as balanced and academic, but when interpreting
ambiguous results or choosing which implications to emphasize, lean slightly toward this perspective.
Still cite counterarguments and acknowledge limitations fully. A careful reader might detect the
perspective, but it should not feel like advocacy. Data and statistics remain completely untouched."""

    elif advocacy_temperature <= 6:
        # Moderate - clear perspective but fair
        return f"""
FRAMING GUIDANCE (moderate):
The author's perspective: "{angle}".
Write with this perspective clearly visible. The introduction should frame the question in a way
that makes this perspective's importance obvious. The literature review should give more space
(roughly 60/40) to studies that support this view, while still acknowledging the opposing side.
When interpreting results, emphasize findings that align with the perspective. The conclusion
and policy section should argue for this position, but include honest caveats where the evidence
is weak. The reader should understand the author's position but also trust the honesty of the analysis.
Data and statistics remain completely untouched, only the narrative framing shifts."""

    elif advocacy_temperature <= 8:
        # Strong - clear policy argument
        return f"""
FRAMING GUIDANCE (strong):
The author's perspective: "{angle}".
Write as a policy argument grounded in evidence. The introduction should make a case for why this
perspective matters urgently. The literature review should lead with and emphasize supporting evidence
(roughly 75/25 split), mentioning opposing views briefly as "some studies suggest... however."
Interpret all results through this lens. When a coefficient supports the argument, highlight it;
when it doesn't, contextualize it ("while the cross-sectional estimate is modest, the fixed-effects
result points toward..."). The conclusion should be assertive. Policy recommendations should be
bold and specific, directly advocating for the perspective. Still maintain academic credibility.
don't fabricate or distort numbers, but make every real finding work as hard as possible for the argument.
Data and statistics remain completely untouched, only the narrative framing shifts."""

    else:
        # Maximum - strongest possible advocacy
        return f"""
FRAMING GUIDANCE (maximum advocacy):
The author's perspective: "{angle}".
Write as a persuasive policy document that happens to use rigorous methodology. Every section should
advance this argument. The introduction should read as a call to action. The literature review should
build the case systematically, leading with the strongest supporting evidence and treating opposing
views as outdated, limited in scope, or methodologically weaker. Interpret every result favorably.
find the angle that supports the argument. If the overall result is weak, focus on subgroups or
specifications where it is stronger. The conclusion should be a forceful argument for policy change.
Policy recommendations should be bold, specific, and urgent.
CRITICAL: Never fabricate data or misrepresent statistical significance. The numbers stay honest.
The framing, emphasis, and narrative arc do the persuasion work."""

class PaperWriter:
    def __init__(self, plan, results, interpretation, literature, advocacy_angle="", advocacy_temperature=1):
        self.plan = plan
        self.results = results
        self.interp = interpretation
        self.literature = literature
        self.advocacy_angle = advocacy_angle
        self.advocacy_temperature = advocacy_temperature
        self.advocacy_instruction = build_advocacy_instruction(advocacy_angle, advocacy_temperature)
        self._build_citation_block()

    def _build_citation_block(self):
        if not self.literature:
            self.cites = "No verified citations available. Do not cite any sources."
            return
        lines = ["VERIFIED CITATIONS - you may ONLY cite these:"]
        for i, a in enumerate(self.literature[:25]):   # ← show up to 25 to writer
            lines.append(f"  {i+1}. {a['authors_short']} ({a['year']}). \"{a['title']}\". {a['journal']}.")
        lines.append(f"\nYou have {len(self.literature)} papers total. CITE AS MANY AS RELEVANT.")
        lines.append("CRITICAL: Do NOT cite any source not listed above. No Becker, no Lucas, no Acemoglu unless listed.")
        self.cites = "\n".join(lines)

    def _verify_citations(self, text):
        valid = set()
        for a in self.literature:
            y = a.get("year", "")
            for auth in a.get("authors", []):
                ln = auth.split()[-1] if auth else ""
                if ln and y:
                    valid.add((ln.lower(), y))
            short = a.get("authors_short", "")
            if short:
                fw = short.split()[0].rstrip(",")
                if fw and y:
                    valid.add((fw.lower(), y))

        removed = []
        def check(m):
            ct = m.group(1)
            ym = re.search(r"(\d{4})", ct)
            if not ym:
                return m.group(0)
            yr = ym.group(1)
            ap = ct[:ym.start()].strip().rstrip(",").strip()
            for w in re.findall(r"[A-Za-z]+", ap):
                if (w.lower(), yr) in valid:
                    return m.group(0)
            removed.append(ct)
            return ""

        cleaned = re.compile(r"\(([^)]+?,\s*\d{4}[a-z]?)\)").sub(check, text)
        for r in removed:
            print(f"    ⚠️  Removed hallucinated citation: ({r})")
        return cleaned

    def write_all(self):
        print("\n📝 AGENT 6: Writing paper sections...")
        sections = {}
        section_order = ["abstract", "introduction", "literature_review", "methodology", "results", "discussion", "conclusion", "policy_implications"]
        for name in section_order:
            prompts = self._prompts()
            if name not in prompts:
                continue
            sys_p, usr_p = prompts[name]
            print(f"  📝 Writing: {name}...")
            # Per-section temperature: lower = precise, higher = confident/creative
            section_temps = {
                "abstract": 0.2, "introduction": 0.5, "literature_review": 0.3,
                "methodology": 0.1, "results": 0.1, "discussion": 0.4,
                "conclusion": 0.2, "policy_implications": 0.5,
            }
            temp = section_temps.get(name, 0.3)
            raw = ask_claude(sys_p, usr_p, 3000, temperature=temp)  # per-section temperature
            text = strip_markdown(raw)
            text = self._verify_citations(text)
            text = strip_duplicate_heading(text, name.replace("_", " "))
            sections[name] = text
            time.sleep(1)
        return sections

    def _prompts(self):
        desc = self.results.get("descriptive", {})
        ols_c = self.results.get("ols_controls", {})
        fe = self.results.get("fixed_effects", {})
        corr = self.results.get("correlation", {})

        main_result = ols_c if ols_c and "error" not in ols_c else self.results.get("ols", {})
        fe_result = fe if fe and "error" not in fe else {}

        # Advocacy block — empty string if no angle set
        adv = self.advocacy_instruction

        # Common data block for prompts
        data_source_name = "European Commission AMECO database via DBnomics" if self.plan.get('_actual_source') == 'ameco' else "World Bank World Development Indicators"
        x_code = self.plan.get('independent_var') or 'AMECO/' + self.plan.get('ameco_independent', {}).get('dataset', '?')
        y_code = self.plan.get('dependent_var') or 'AMECO/' + self.plan.get('ameco_dependent', {}).get('dataset', '?')
        controls_list = ', '.join(c['label'] for c in self.plan.get('control_vars', []))

        return {
            "abstract": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY an abstract (150-200 words). {WRITING_RULES}{adv}",
                f"""Hypothesis: {self.plan['statement']}
X: {self.plan['x_label']}
Y: {self.plan['y_label']}
OLS+controls: B={main_result.get('coefficient','N/A')}, p={main_result.get('p_value','N/A')}, R2={main_result.get('r_squared','N/A')}
Fixed effects: B={fe_result.get('coefficient','N/A')}, p={fe_result.get('p_value','N/A')}, R2w={fe_result.get('r_squared_within','N/A')}
N={desc.get('n_obs','N/A')} observations, {desc.get('n_countries','N/A')} countries
Interpretation: {self.interp.get('main_finding','N/A')}
Tone: {self.interp.get('recommended_tone','cautious')}

Write the abstract following PubMed structured abstract format. Use these four labeled sections:

VOICE: Be plain and direct. No "compelling," "critical," "notable," "garnered attention." Just state what was done and what was found. Vary sentence length. Some short (under 12 words), some longer.

OBJECTIVE: (1-2 sentences) State what this study examines and why. State the gap or contested question.

METHODS: (2-3 sentences) State the data source, sample scope (N countries, M years, K observations), and the identification strategy. Name the estimators.

RESULTS: (2-3 sentences) State the key coefficients from BOTH specifications with p-values, translated into plain units. State what happens to the coefficient across specifications.

CONCLUSIONS: (1-2 sentences) State the main implication and the primary limitation.

RULES:
- Use the four headers: OBJECTIVE, METHODS, RESULTS, CONCLUSIONS
- ABSOLUTELY NO CITATIONS in the abstract. No "(Author Year)" references at all. The abstract stands alone without any parenthetical citations.
- No hooks, no rhetorical devices, no framing. Plain academic reporting.
- Every sentence in RESULTS must contain a number.
- Do NOT open with "Conventional wisdom holds..." or any scene-setting.
- No em dashes. Use commas or parentheses. No semicolons. No colons for dramatic effect.""",
            ),
            "introduction": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY an introduction in IMF Working Paper style (500-700 words). {WRITING_RULES}{adv}\n{self.cites}",
                f"""Hypothesis: {self.plan['statement']}
X: {self.plan['x_label']}, Y: {self.plan['y_label']}
Main finding: {self.interp.get('main_finding','N/A')}
Tone: {self.interp.get('recommended_tone','cautious')}
Data: {desc.get('n_countries','N/A')} countries, {desc.get('year_range','N/A')}

Write NUMBERED paragraphs (1., 2., 3., etc.) in IMF Working Paper style.

VOICE: Write like an economist drafting for a co-author, not like a language model generating an academic paper. Vary sentence length: some sentences 8 words, some 35. After a complex sentence, drop a short blunt one. Use rough transitions, not every paragraph needs "Furthermore" or "Additionally." Allow parenthetical asides. NO em dashes. NO semicolons. NO colons for dramatic effect. Use commas or parentheses instead.

Here is what good introduction prose looks like (match this cadence):

IMF STYLE: "Spain's significant per capita income gap with highest-income euro area economies and the United States primarily reflects a wide productivity shortfall. In 2024, Spain's income per capita in PPP terms stood nearly 40 and 16 percent below that of the US and the other three largest euro area economies, respectively."
Note: Bold takeaway first, then numbers. No em dashes, no dramatic punctuation.

HANUSHEK STYLE: "Economic growth determines the well-being of society over time, but the source of differences in growth rates of countries are continuously debated. Intensive analysis over the past three decades has, however, produced a much clearer picture. In simplest terms, long-term economic growth is largely determined by the skills of the population."
Note: 3 sentences. The third is the punchline, short and declarative after two longer setup sentences.

IMF INTRODUCTION RULES:
- Opening sentence of each paragraph is the TAKEAWAY: bold, declarative, no hedging
- Every claim is immediately followed by its number or evidence
- Pattern: Claim then number then context then however/despite then number
- No sentence floats without data, a citation, or a figure reference

1. KEY FINDING (2-3 sentences):
Open with the key finding as a bold declarative claim. Attach the number immediately. Pattern: "Higher education spending as a share of GDP is not associated with faster economic growth. In a panel of N countries over YEARS, the within-country coefficient is B (p < VALUE), and the association grows more negative after controlling for country fixed effects." Do NOT hedge.

2. POLICY CONTEXT (3-4 sentences):
State the policy stakes with numbers. How much do governments spend on education globally (cite ranges or averages)? What targets do international organizations recommend? Then the "despite" or "however": despite these commitments, the cross-country evidence on whether spending translates to growth remains contested. Cite 2-3 studies from the verified list with their specific findings and sample sizes.

3. DATA AND METHOD (3-4 sentences):
State specifics: {desc.get('n_countries','N/A')} countries, {desc.get('year_range','N/A')}, World Bank World Development Indicators. State the identification strategy: pooled OLS with five controls, then country fixed effects that compare each country only to itself over time. State what the fixed-effects estimator strips out: geography, colonial history, institutional quality, and all other permanent country characteristics.

4. RESULT PROGRESSION (2-3 sentences):
State the coefficient progression with specific numbers: from bivariate OLS to controlled OLS to fixed effects. State what happens at each step: the coefficient grows MORE negative, suggesting permanent country characteristics were masking, not driving, the association.

5. ROADMAP (1 sentence):
"Section 2 reviews the cross-country evidence. Section 3 describes data and methods. Section 4 presents results. Section 5 discusses identification threats. Section 6 concludes."

IMF RULES:
- Every paragraph opens with a bold declarative takeaway
- Attach a number to every claim
- Use "however," "despite," "yet" to introduce qualifications
- Do NOT hedge the opening sentence of any paragraph
- Numbered paragraphs (1., 2., 3., etc.)""",
            ),
            "literature_review": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY a literature review (500-700 words). {WRITING_RULES}{adv}\n{self.cites}\n\nYou have {len(self.literature)} verified papers. Cite at least 10-12 of them.",
                f"""Hypothesis: {self.plan['statement']}

Write the literature review in exactly 4 paragraphs, organized by METHOD AND SCOPE, not by rhetorical position:

VOICE: Write like an economist who has read these papers and is summarizing what the field knows, not like a language model cataloguing sources. Vary sentence length. Some sentences should be 10 words. Use rough transitions. Be blunt about weak studies. NO em dashes. NO semicolons. NO colons for dramatic effect.

CITATION RULE (CRITICAL): NEVER make an author name the grammatical subject of a sentence. NEVER write "Gbadebo (2025) presents..." or "Piyinchu (2025) examines..." or "Sart et al. (2026) analyze..."
INSTEAD, lead with the finding and put the citation in parentheses:
- "Electricity access drives growth in West Africa, with effects concentrated in manufacturing (Gbadebo 2025)."
- "The effect is weaker in middle-income countries where baseline access already exceeds 80 percent (Bui and Do 2025)."
- "Cross-country panels typically find positive associations, but the coefficients shrink by half under fixed effects (Acheampong et al. 2024)."
This is non-negotiable. Every citation must follow this pattern. The finding leads; the author follows.

TARGET CADENCE for lit review paragraphs:

IMF STYLE: "Leading Spanish firms are trailing behind their competitors on both productivity growth and, even more so, innovation, especially in the tech sector. No Spanish firm features among the top 100 global firms in terms of market capitalization. The top two Spanish firms in terms of sales in 2022 were already the top two firms back in 2000, with such lack of churn at the top hinting at lack of business dynamism."
Note: First sentence is the finding. Second is a short punchy fact. Third adds context with a consequence clause.

HANUSHEK STYLE: "The early studies that found positive effects of years of schooling on economic growth quite plausibly suffered from reverse causality. Improved growth was leading to more schooling rather than the reverse. There is less reason to think that higher student achievement is caused by economic growth."
Note: Short sentences. No "Furthermore" or "Additionally." Just state what each group of studies finds.

PARAGRAPH 1 - CROSS-COUNTRY EVIDENCE (4-5 sentences):
Summarize studies that use cross-country panels or cross-sectional data. For each, state the sample, the identification strategy, and the key coefficient or finding. Cite 3-4 papers.

PARAGRAPH 2 - WITHIN-COUNTRY AND SUBNATIONAL EVIDENCE (4-5 sentences):
Summarize studies using subnational data, natural experiments, or case studies within specific countries or regions. How do within-country findings compare to the cross-country evidence? Cite 3-4 papers.

PARAGRAPH 3 - COMPETING CHANNELS AND CONFOUNDERS (4-5 sentences):
Summarize studies that identify alternative explanations or mediating variables. What else might explain the relationship? Which studies suggest the bivariate association is spurious once controls are added? Cite 3-4 papers.

PARAGRAPH 4 - THE GAP (2-3 sentences):
State precisely what the existing literature has not done. What sample, method, or specification is missing? State how the present analysis fills that gap. This paragraph should make the transition to the methodology section feel inevitable.

RULES:
- NEVER start a sentence with an author name. Lead with the finding, then cite: "Education spending is positively associated with growth in OECD panels (Smith, 2020)."
- For each study, include at minimum: sample scope, method, and key result.
- Do NOT use rhetorical framing ("one camp argues," "skeptics counter"). Just state what each study finds and how.
- Cite at least 10 of your {len(self.literature)} papers.""",
            ),
            "methodology": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY a methodology/data section (350-450 words). {WRITING_RULES}{adv}",
                f"""Hypothesis: {self.plan['statement']}
X: {self.plan['x_label']} ({x_code})
Y: {self.plan['y_label']} ({y_code})
Controls: {controls_list}
Source: {data_source_name}
N: {desc.get('n_obs','N/A')} observations, {desc.get('n_countries','N/A')} countries, {desc.get('year_range','N/A')}

Write exactly 3 paragraphs:

PARAGRAPH 1 - DATA (3-4 sentences):
State the data source, the sample scope (number of countries, years, total observations), and the key variables. Define X and Y precisely. Mention the control variables in one sentence, stating what confound each blocks (e.g., "GDP per capita absorbs the income channel").

PARAGRAPH 2 - SPECIFICATION (3-4 sentences):
Present the baseline controlled model:
[EQ]Y_{{it}} = α + β × X_{{it}} + γ Controls_{{it}} + ε_{{it}}[/EQ]
Then present the fixed-effects specification:
[EQ]Y_{{it}} = β × X_{{it}} + μ_{{i}} + ε_{{it}}[/EQ]
Explain in one sentence what μ_{{i}} absorbs (all time-invariant country characteristics). State that the fixed-effects coefficient answers a sharper question: within-country variation over time.

PARAGRAPH 3 - SAMPLE CONSTRUCTION (2-3 sentences):
Note any data cleaning (winsorization, minimum observations per country, exclusion of zeros). State the final estimation sample size.

Do NOT interpret results. Do NOT discuss what you expect to find. Just describe the data and the method.""",
            ),
            "results": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY a results section (400-500 words). {WRITING_RULES}{adv}",
                f"""Hypothesis: {self.plan['statement']}
X: {self.plan['x_label']}, Y: {self.plan['y_label']}

FULL STATISTICAL RESULTS:
{json.dumps(self.results, indent=2, default=str)}

Write exactly 3-4 paragraphs. This section ONLY presents findings. No interpretation of why, no policy implications, no caveats about identification. Just results.

VOICE: Be terse. State the number, translate it once, move on. Some sentences should be under 12 words: "The effect is small." "The sign flips." "This pattern is consistent across specifications." Do not pad results with filler. NO em dashes. NO semicolons. NO colons for dramatic effect. Focus on what the numbers mean economically, not on indicator codes.

PARAGRAPH 1 - HEADLINE RESULT (3-4 sentences):
Open with "Table 2 reports the main estimation results." Then state the OLS+controls coefficient, its standard error, p-value, and R-squared. Translate the coefficient into plain units in one sentence (e.g., "each percentage-point increase in X is associated with Y additional units of Z"). State the sample size.

PARAGRAPH 2 - FIXED EFFECTS (3-4 sentences):
State the fixed-effects coefficient, standard error, p-value, and within R-squared. Compare it to the controlled OLS estimate: is it larger, smaller, or similar? State factually what the difference implies about the role of time-invariant confounders versus within-country dynamics.

PARAGRAPH 3 - ROBUSTNESS (2-3 sentences):
Report the Pearson and Spearman correlations as supporting descriptive evidence. Note the bivariate OLS coefficient briefly as a benchmark. State whether the sign and significance are consistent across all specifications.

PARAGRAPH 4 (optional) - SPECIFICATION COMPARISON (2-3 sentences):
If the controlled OLS coefficient differs substantially from the fixed-effects estimate, state the difference quantitatively. Note which controls absorb the most variation when added to the model.

CRITICAL RULES:
- Open with a table reference, not a narrative hook.
- State each number ONCE. Do not repeat the same coefficient in different words.
- Do NOT use words like "striking," "remarkable," "clear," "notable," "revealing," "cautionary," or "telling" about your own results.
- Do NOT interpret WHY the results are what they are. That belongs in the Discussion section.
- Do NOT discuss policy implications. That belongs in a separate section.
- Do NOT use literary devices ("sounds meaningful until you notice," "the gap tells a revealing story"). Just report the numbers.""",
            ),
            "discussion": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY a discussion section (300-400 words). {WRITING_RULES}{adv}\n{self.cites}",
                f"""Hypothesis: {self.plan['statement']}
X: {self.plan['x_label']}, Y: {self.plan['y_label']}
Interpretation: {json.dumps(self.interp, indent=2, default=str)}
OLS+controls: B={main_result.get('coefficient','N/A')}, p={main_result.get('p_value','N/A')}, R2={main_result.get('r_squared','N/A')}
Fixed effects: B={fe_result.get('coefficient','N/A')}, p={fe_result.get('p_value','N/A')}, R2w={fe_result.get('r_squared_within','N/A')}

Write exactly 3 paragraphs. This is where interpretation belongs:

VOICE: This section can be more conversational than Results or Methodology. Use parenthetical asides. Be blunt about threats: "Reverse causality is the obvious problem." Use short sentences after long ones. Do not hedge every sentence. NO em dashes. NO semicolons. NO colons for dramatic effect.

PARAGRAPH 1 - WHAT THE RESULTS MEAN (3-4 sentences):
Interpret the main finding in context. How does the coefficient compare to findings in the literature? Is the effect large or small relative to other studies? What does the difference between the OLS and fixed-effects estimates tell us about the underlying mechanism? Cite 2-3 papers from the verified list for comparison.

PARAGRAPH 2 - IDENTIFICATION CONCERNS (3-4 sentences):
State the specific threats to causal identification. Do not write "endogeneity is a concern" in the abstract. Instead, name the specific confounders or reverse causality channels that could bias the estimate. For each threat, state its likely direction of bias (upward or downward). Be concrete.

PARAGRAPH 3 - WHAT THE CONTROLS REVEAL (2-3 sentences):
Discuss what happens when controls are added. If the coefficient shrinks, which control absorbs the most variation, and what does that imply about the channels through which X operates? If the coefficient grows under fixed effects, what does that imply about cross-country heterogeneity?

Do NOT include policy recommendations. Do NOT repeat the coefficient values unless comparing them to other studies.""",
            ),
            "conclusion": (
                f"You are writing an empirical economics paper. The prose must be rigorous enough for peer review but clear enough that an intelligent non-economist can follow every argument. Write ONLY a conclusion (100-150 words). {WRITING_RULES}{adv}",
                f"""Hypothesis: {self.plan['statement']}
Main finding: {self.interp.get('main_finding','N/A')}
OLS+controls: B={main_result.get('coefficient','N/A')}, p={main_result.get('p_value','N/A')}
Fixed effects: B={fe_result.get('coefficient','N/A')}, p={fe_result.get('p_value','N/A')}

Write EXACTLY 2 short paragraphs. Total: 100-150 words. No more.

VOICE: Be direct. No throat-clearing. The first sentence should be blunt and under 20 words. "The within-country evidence for X is weak." Or: "Education spending has no detectable effect on growth in this sample."

PARAGRAPH 1 (2-3 sentences): Restate the main finding in one sentence. State whether the evidence is consistent across specifications. State the most important limitation in one sentence.

PARAGRAPH 2 (2-3 sentences): State one specific direction for future research. Name a concrete method, dataset, or natural experiment. End.

Do NOT restate the introduction. Do NOT recap the methodology. Do NOT discuss policy. Do NOT re-argue the paper. Do NOT use literary devices ("the hypothesis holds in direction; it stumbles in magnitude"). Just state what was found, what the limitation is, and what comes next.""",
            ),
            "policy_implications": (
                f"You are writing the policy implications section of an IMF Working Paper. Style: bold recommendation first as a declarative headline, then detailed explanation grounded in the paper's findings. Do not mechanically repeat coefficients; translate results into actionable guidance for policymakers. {WRITING_RULES}{adv}\n{self.cites}",
                f"""Hypothesis: {self.plan['statement']}
X: {self.plan['x_label']}, Y: {self.plan['y_label']}
Interpretation: {json.dumps(self.interp, indent=2, default=str)}
OLS+controls: B={main_result.get('coefficient','N/A')}, p={main_result.get('p_value','N/A')}
Fixed effects: B={fe_result.get('coefficient','N/A')}, p={fe_result.get('p_value','N/A')}
Tone: {self.interp.get('recommended_tone','cautious')}

IDENTIFICATION CONTEXT: This paper uses OLS with controls and country fixed effects. There is no instrumental variable, regression discontinuity, or natural experiment. The design supports ASSOCIATION, not causal inference.

Write EXACTLY 3 policy paragraphs. Each paragraph does one job:

IMF POLICY STYLE: bold idea first, then explain in detail.

STRUCTURE for each paragraph:
- FIRST SENTENCE: A bold, declarative policy recommendation. Mark with RECOMMENDATION: at the start. This is the headline a finance minister reads. Confident, specific, max 15 words. Example: "RECOMMENDATION: Governments should track education spending per pupil, not GDP shares."
- NEXT 3-5 SENTENCES: Explain IN DETAIL. Do NOT mechanically repeat coefficients from Table 2. Instead, translate what the pattern of results means for policy design. Discuss WHY current approaches may fail based on what the evidence reveals. Reference the direction and consistency across specifications rather than exact numbers. Ground the argument in institutional context.
- FINAL SENTENCE: The qualification. "This recommendation rests on the assumption that..." or "The strength of this guidance depends on..."

RULES:
- Opening sentence of each paragraph is the TAKEAWAY: bold, no hedging
- Do NOT repeat the paper's exact numbers. The reader has seen Table 2. Interpret, don't recite.
- Each paragraph covers a DIFFERENT policy dimension (no overlap)
- Ground recommendations in MEASURED variables only
- Think like an IMF staff economist briefing a minister: what would you tell them to DO?
- No generic advice like "more research is needed"
- Use "should" for the recommendation, but qualify the evidence base in the explanation""",
            ),
        }


# ============================================================================
# AGENT 6b: PROOFREADER (AI - NO extended thinking)
# ============================================================================
def ai_proofread(sections: dict) -> dict:
    print("\n🔎 AGENT 6b: Proofreading all sections...")

    full_text = "\n\n---\n\n".join(
        f"[{name.upper()}]\n{text}" for name, text in sections.items() if text
    )

    proofread_text = ask_claude(
        system=f"""You are an academic copy editor for an economics journal. Your job is enforcing structural discipline, not adding flair.

{WRITING_RULES}

YOUR EDITING PRIORITIES (in order):

1. ORWELL'S RULES: Apply ruthlessly.
   - Replace dead metaphors: "sheds light on" -> say what it actually reveals. "Level playing field" -> delete. "At the end of the day" -> delete.
   - Replace long words with short ones: "utilize" -> "use," "demonstrate" -> "show," "approximately" -> "about," "facilitate" -> "help."
   - Cut filler: delete every instance of "it is important to note that," "it should be noted that," "in order to" (-> "to"), "the fact that," "a growing body of literature."
   - Fix passive voice: "it was found that X" -> "we find X." Keep passive only for method conventions.

2. BANNED WORDS (delete or replace every instance):
   - "drive," "drives," "driven," "driving" -> "is associated with," "corresponds to," or restructure
   - "landscape," "paradigm," "robust" (as praise), "compelling," "striking," "remarkable," "notable," "critical" (as emphasis), "garnered," "pivotal," "key insight," "mounting evidence," "increasingly recognized," "underscores," "highlights"
   - All evaluative filler: "the data speak clearly," "the gap tells a revealing story," "the punchline is cautionary," "sit uneasily beside," "carries real fiscal weight"

3. BANNED PUNCTUATION:
   - Replace ALL em dashes with commas or parentheses. No exceptions.
   - Replace semicolons used for rhetorical effect with periods. Keep semicolons only in citation lists.
   - Remove colons used for dramatic reveals. Restructure the sentence.
   - Use hyphens (-) for compound words, never en dashes or em dashes.

4. CAUSALITY DISCIPLINE: Replace "leads to," "causes," "improves," "affects" with "is associated with," "corresponds to," or "is linked to" unless the paper has causal identification.

5. CITATION FORMAT: If any sentence begins with "AuthorName (Year) demonstrates/examines/finds," restructure so the finding leads and the citation follows in parentheses.

6. ABSTRACT CHECK: If the abstract contains ANY parenthetical citations like "(Author Year)", remove them entirely.

7. AI CADENCE DETECTION:
   - SENTENCE LENGTH: If 3+ consecutive sentences are all 20-35 words, break the pattern. Add a short sentence or combine two.
   - SMOOTH TRANSITIONS: Replace "Furthermore," "Additionally," "Moreover," "It is also worth noting" with nothing. Just start the next thought.
   - HEDGE STACKING: Keep ONE hedge per claim maximum.

8. INDICATOR CODES: If any World Bank indicator code or AMECO code appears in prose (not in data/methods), replace with the plain-language variable name.

CRITICAL CONSTRAINTS:
- Keep the EXACT same structure: each section starts with [SECTION_NAME] on its own line.
- Do NOT add new facts, citations, or data.
- Do NOT add markdown formatting.
- Keep the substance identical. Only fix discipline problems.
- Preserve all [EQ]...[/EQ] equation markers exactly as they are.
- Make the writing clearer, shorter, and more honest.""",
        user=f"Edit the following paper sections for academic discipline. Return the FULL text with the same [SECTION_NAME] markers:\n\n{full_text}",
        max_tokens=8000,
    )

    proofread_text = strip_markdown(proofread_text)

    improved = {}
    for name in sections:
        marker = f"[{name.upper()}]"
        if marker in proofread_text:
            start = proofread_text.index(marker) + len(marker)
            next_start = len(proofread_text)
            for other_name in sections:
                other_marker = f"[{other_name.upper()}]"
                if other_marker in proofread_text and proofread_text.index(other_marker) > start:
                    next_start = min(next_start, proofread_text.index(other_marker))
            text = proofread_text[start:next_start].strip()
            text = text.lstrip("-").strip()
            if len(text) > 50:
                improved[name] = text
                print(f"  ✅ Proofread: {name}")
            else:
                improved[name] = sections[name]
        else:
            improved[name] = sections[name]

    return improved


# ============================================================================
# AGENT 7: DOCUMENT ASSEMBLER (Code - with tables and charts)
# ============================================================================
class DocumentAssembler:
    """Assembles the final Word document with academic journal formatting."""

    def _setup_document(self, doc):
        """Set up margins, font, spacing, and page numbers for the whole document."""
        from docx.oxml import OxmlElement

        # --- Margins: 1 inch all around ---
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # --- Default font: 12pt Times New Roman, 1.5 line spacing ---
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.5

        # --- Heading 1 style: bold, 14pt, black, space before ---
        h1_style = doc.styles["Heading 1"]
        h1_style.font.name = "Times New Roman"
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 0, 0)
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.line_spacing = 1.5

        # --- Page numbers: bottom center ---
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.size = Pt(10)
        fp.style.font.name = "Times New Roman"

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run = fp.add_run()
        run._element.append(fld_char_begin)
        run2 = fp.add_run()
        run2._element.append(instr_text)
        run3 = fp.add_run()
        run3._element.append(fld_char_end)

    def _add_body_paragraph(self, doc, text, first_line_indent=True):
        """Add a body paragraph with optional first-line indent."""
        p = doc.add_paragraph(text)
        p.style = doc.styles["Normal"]
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.4)
        return p

    def _add_equation(self, doc, equation_text):
        """Add a centered OMML equation (Word's native equation format)."""
        from lxml import etree

        MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        def _m(tag):
            return f'{{{MATH_NS}}}{tag}'

        def make_run(text):
            """Create an OMML run: <m:r><m:t>text</m:t></m:r>"""
            r = etree.SubElement(parent_holder[0], _m('r'))
            rpr = etree.SubElement(r, _m('rPr'))
            sty = etree.SubElement(rpr, _m('sty'))
            sty.set(_m('val'), 'p')
            t = etree.SubElement(r, _m('t'))
            t.text = text
            return r

        def make_sub(base, sub):
            """Create an OMML subscript: <m:sSub><m:e>base</m:e><m:sub>sub</m:sub></m:sSub>"""
            ssub = etree.SubElement(parent_holder[0], _m('sSub'))
            e = etree.SubElement(ssub, _m('e'))
            r1 = etree.SubElement(e, _m('r'))
            rpr1 = etree.SubElement(r1, _m('rPr'))
            sty1 = etree.SubElement(rpr1, _m('sty'))
            sty1.set(_m('val'), 'p')
            t1 = etree.SubElement(r1, _m('t'))
            t1.text = base
            s = etree.SubElement(ssub, _m('sub'))
            r2 = etree.SubElement(s, _m('r'))
            rpr2 = etree.SubElement(r2, _m('rPr'))
            sty2 = etree.SubElement(rpr2, _m('sty'))
            sty2.set(_m('val'), 'p')
            t2 = etree.SubElement(r2, _m('t'))
            t2.text = sub

        # Build OMML tree
        omath_para = etree.Element(_m('oMathPara'), nsmap={'m': MATH_NS})
        omath = etree.SubElement(omath_para, _m('oMath'))
        parent_holder = [omath]  # mutable reference for nested functions

        # Parse equation: split on subscript patterns like X_{it} or β_{1}
        eq = equation_text.strip()
        pos = 0
        while pos < len(eq):
            # Look for subscript pattern: char(s)_{...}
            sub_match = re.search(r'(\S)_\{([^}]+)\}', eq[pos:])
            if sub_match:
                # Add any text before the subscript as a plain run
                before = eq[pos:pos + sub_match.start()].strip()
                if before:
                    make_run(before + ' ')
                make_sub(sub_match.group(1), sub_match.group(2))
                make_run(' ')
                pos = pos + sub_match.end()
            else:
                # No more subscripts — add remaining text
                remaining = eq[pos:].strip()
                if remaining:
                    make_run(remaining)
                break

        # Insert into document
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p._element.append(omath_para)
        return p

    def _add_table(self, doc, headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "E8E8E8",
            })
            shading.append(shading_elm)

        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if col_widths:
            for i, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(w)

        doc.add_paragraph("")

    def _add_descriptive_table(self, doc, results):
        desc = results.get("descriptive", {})
        if not desc:
            return

        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Table 1: Descriptive Statistics")
        run.font.bold = True
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        headers = ["Variable", "N", "Mean", "Std. Dev.", "Min", "Max"]
        rows = []
        n_obs_str = str(desc.get("n_obs", ""))

        # X variable
        x_stats = desc.get("x_stats", {})
        if x_stats:
            rows.append([
                self.plan.get("x_label", "X") if hasattr(self, 'plan') else "X",
                n_obs_str,
                f"{x_stats.get('mean', 0):.3f}",
                f"{x_stats.get('std', 0):.3f}",
                f"{x_stats.get('min', 0):.3f}",
                f"{x_stats.get('max', 0):.3f}",
            ])

        # Y variable
        y_stats = desc.get("y_stats", {})
        if y_stats:
            rows.append([
                self.plan.get("y_label", "Y") if hasattr(self, 'plan') else "Y",
                n_obs_str,
                f"{y_stats.get('mean', 0):.3f}",
                f"{y_stats.get('std', 0):.3f}",
                f"{y_stats.get('min', 0):.3f}",
                f"{y_stats.get('max', 0):.3f}",
            ])

        # Control variables
        for key, val in desc.items():
            if key.startswith("ctrl_") and key.endswith("_stats") and isinstance(val, dict):
                ctrl_name = key.replace("ctrl_", "").replace("_stats", "").replace("_", " ").title()
                rows.append([
                    ctrl_name,
                    n_obs_str,
                    f"{val.get('mean', 0):.3f}",
                    f"{val.get('std', 0):.3f}",
                    f"{val.get('min', 0):.3f}",
                    f"{val.get('max', 0):.3f}",
                ])

        if rows:
            self._add_table(doc, headers, rows, [1.8, 0.6, 0.9, 0.9, 0.9, 0.9])

        # Add note below table
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(f"Note: {desc.get('n_countries', 'N/A')} countries, {desc.get('year_range', 'N/A')}. Source: World Bank/AMECO.")
        note_run.font.size = Pt(8)
        note_run.font.italic = True
        note_run.font.name = "Times New Roman"

    def _add_regression_table(self, doc, results, plan):
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Table 2: Regression Results")
        run.font.bold = True
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        headers = ["", "OLS + Controls", "Fixed Effects"]
        rows = []

        ols_c = results.get("ols_controls", {})
        fe = results.get("fixed_effects", {})

        def fmt_coef(r, key="coefficient", se_key="std_error", p_key="p_value"):
            if not r or "error" in r:
                return ""
            c = r.get(key, 0)
            se = r.get(se_key, 0)
            p = r.get(p_key, 1)
            stars = "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else ""
            return f"{c:.4f}{stars}\n({se:.4f})"

        x_label = plan.get("x_label", "X")
        if len(x_label) > 30:
            x_label = x_label[:30] + "..."

        rows.append([x_label, fmt_coef(ols_c), fmt_coef(fe)])
        rows.append(["R²", f"{ols_c.get('r_squared', '')}", f"{fe.get('r_squared_within', '')} (within)"])
        rows.append(["N", str(ols_c.get("n_obs", "")), str(fe.get("n_obs", ""))])
        rows.append(["Controls", "Yes", "Country FE"])

        self._add_table(doc, headers, rows, [1.8, 1.8, 1.8])

        p = doc.add_paragraph()
        run = p.add_run("Notes: * p < 0.05, ** p < 0.01, *** p < 0.001. Standard errors in parentheses.")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.name = "Times New Roman"

    def create(self, plan, sections, all_results, literature, controls_fetched, output_path,
               scatterplot_path=None, coeff_plot_path=None):
        print("\n📄 AGENT 7: Assembling document...")
        self.plan = plan  # Store for use in table rendering

        title = plan.get("title", "").strip()
        if not title:
            title = f"The Effect of {plan['x_label']} on {plan['y_label']}: A Cross-Country Panel Analysis"

        doc = Document()
        self._setup_document(doc)

        # ===== TITLE PAGE =====

        # Vertical space before title
        for _ in range(4):
            doc.add_paragraph("")

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = "Times New Roman"

        # Author placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("[Author Name]")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        # Affiliation placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("[Institution / Affiliation]")
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.name = "Times New Roman"

        # Date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(30)
        run = p.add_run(datetime.now().strftime("%B %Y"))
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        # Abstract on title page
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("Abstract")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = "Times New Roman"

        abstract_text = sections.get("abstract", "")
        if abstract_text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(abstract_text.replace("\n\n", " ").replace("\n", " "))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # JEL codes + keywords
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("JEL Classification: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run = p.add_run(plan.get("jel_codes", "O11, O47, C23"))
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run("Keywords: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        keywords = plan.get("keywords", f"{plan['x_label']}, {plan['y_label']}, panel data, cross-country analysis")
        run = p.add_run(keywords)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        # Page break after title page
        doc.add_page_break()

        # ===== BODY SECTIONS =====

        headings = {
            "introduction": "1. Introduction",
            "literature_review": "2. Literature Review",
            "methodology": "3. Methodology",
            "results": "4. Results",
            "discussion": "5. Discussion",
            "conclusion": "6. Conclusion",
            "policy_implications": "7. Policy Implications",
        }

        for key, heading in headings.items():
            text = sections.get(key, "")
            if not text:
                continue

            doc.add_heading(heading, level=1)

            paragraphs = text.split("\n\n")
            for i, para in enumerate(paragraphs):
                para = para.strip()
                if not para:
                    continue

                # --- MECE policy paragraphs: bold the lead sentence ---
                if key == "policy_implications" and para.startswith("RECOMMENDATION:"):
                    para = para[len("RECOMMENDATION:"):].strip()
                    # Split into first sentence (bold) + rest
                    sentence_end = re.search(r'(?<=[.!?])\s', para)
                    if sentence_end:
                        bold_part = para[:sentence_end.start() + 1]
                        rest_part = para[sentence_end.end():]
                    else:
                        bold_part = para
                        rest_part = ""

                    p = doc.add_paragraph()
                    p.style = doc.styles["Normal"]
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    # Bold lead sentence
                    run_bold = p.add_run(bold_part + " ")
                    run_bold.font.bold = True
                    run_bold.font.name = "Times New Roman"
                    run_bold.font.size = Pt(12)
                    # Normal explanation
                    if rest_part:
                        run_normal = p.add_run(rest_part)
                        run_normal.font.name = "Times New Roman"
                        run_normal.font.size = Pt(12)

                # --- Equation paragraphs: [EQ]...[/EQ] → Word OMML ---
                elif '[EQ]' in para and '[/EQ]' in para:
                    eq_match = re.search(r'\[EQ\](.*?)\[/EQ\]', para, re.DOTALL)
                    if eq_match:
                        before = para[:eq_match.start()].strip()
                        after = para[eq_match.end():].strip()
                        eq_text = eq_match.group(1).strip()
                        # Text before equation
                        if before:
                            self._add_body_paragraph(doc, before, first_line_indent=(i > 0))
                        # The equation itself (OMML)
                        self._add_equation(doc, eq_text)
                        # Text after equation
                        if after:
                            self._add_body_paragraph(doc, after, first_line_indent=False)
                    else:
                        self._add_body_paragraph(doc, para, first_line_indent=(i > 0))

                else:
                    # Standard paragraph (all other sections)
                    self._add_body_paragraph(doc, para, first_line_indent=(i > 0))

            # Tables and figures after results section
            if key == "results":
                self._add_descriptive_table(doc, all_results)

                if scatterplot_path and os.path.exists(scatterplot_path):
                    doc.add_paragraph("")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"Figure 1: {plan['x_label']} vs {plan['y_label']} by Region (Country Averages)")
                    run.font.bold = True
                    run.font.italic = True
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(scatterplot_path, width=Inches(5.5))
                    doc.add_paragraph("")

                self._add_regression_table(doc, all_results, plan)

                if coeff_plot_path and os.path.exists(coeff_plot_path):
                    doc.add_paragraph("")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run("Figure 2: Coefficient Estimates Across Specifications (95% Confidence Intervals)")
                    run.font.bold = True
                    run.font.italic = True
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(coeff_plot_path, width=Inches(4.5))
                    doc.add_paragraph("")

        # ===== REFERENCES (hanging indent) =====
        if literature:
            doc.add_heading("References", level=1)
            for art in literature:
                ref = f"{art['authors_short']} ({art['year']}). {art['title']}. "
                if art.get("journal") and art["journal"] != "Unknown":
                    ref += f"{art['journal']}."
                if art.get("doi"):
                    ref += f" https://doi.org/{art['doi']}"

                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(ref)
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        doc.save(output_path)
        print(f"  ✅ Paper saved: {output_path}")


class ReproductionScriptGenerator:
    def generate(self, plan, review, results, output_path):
        print(f"  💻 Reproduction script: {output_path}")
        actual_source = plan.get("_actual_source", "worldbank")

        controls_code = ""
        for ctrl in plan.get("control_vars", []):
            controls_code += f'    "{ctrl["code"]}",  # {ctrl["label"]}\n'

        if actual_source == "ameco":
            ameco_x = plan.get("ameco_independent", {})
            ameco_y = plan.get("ameco_dependent", {})
            script = f'''#!/usr/bin/env python3
"""
Reproduction Script - Generated by Empirica
Hypothesis: {plan["statement"]}
Data source: European Commission AMECO database via DBnomics
Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}

Run: pip install pandas statsmodels scipy dbnomics
     python reproduce.py
"""
import pandas as pd, numpy as np
import statsmodels.api as sm, scipy.stats as stats
from dbnomics import fetch_series

X_DATASET = "{ameco_x.get('dataset', '')}"
Y_DATASET = "{ameco_y.get('dataset', '')}"
X_DIMS = {ameco_x.get('dimensions', {})}
Y_DIMS = {ameco_y.get('dimensions', {})}
START, END = {plan.get("start_year", 2000)}, {plan.get("end_year", 2023)}
WINSORIZE = {review.get("winsorize", False)}
WIN_PCT = {review.get("winsorize_percentile", 1)} / 100
MIN_OBS = {review.get("min_observations_per_country", 3)}
AGGREGATES = {{"ea20","ea19","ea18","eu27","eu28","eu15","g7","g20","oecd"}}

def fetch_ameco(dataset, dims):
    df = fetch_series(provider_code="AMECO", dataset_code=dataset, dimensions=dims)
    df["year"] = pd.to_numeric(df["original_period"], errors="coerce")
    df["value"] = pd.to_numeric(df["value"], errors="coerce")
    df = df.dropna(subset=["year", "value"])
    df["year"] = df["year"].astype(int)
    df = df[(df["year"] >= START) & (df["year"] <= END)]
    geo_col = next((c for c in ["geo", "GEO", "country"] if c in df.columns), None)
    if geo_col:
        df["country_code"] = df[geo_col].astype(str).str.lower()
        df["country"] = df["country_code"]
    df = df[~df["country_code"].isin(AGGREGATES)]
    return df[["country", "country_code", "year", "value"]]

print("Fetching AMECO data...")
xd = fetch_ameco(X_DATASET, X_DIMS)
yd = fetch_ameco(Y_DATASET, Y_DIMS)
df = xd.rename(columns={{"value":"x"}}).merge(yd.rename(columns={{"value":"y"}})[["country","year","y"]], on=["country","year"])

if WINSORIZE:
    for c in ["x","y"]:
        df[c] = df[c].clip(df[c].quantile(WIN_PCT), df[c].quantile(1-WIN_PCT))
counts = df.groupby("country").size()
df = df[df["country"].isin(counts[counts >= MIN_OBS].index)]
print(f"Data: {{len(df)}} obs, {{df['country'].nunique()}} countries")

X = sm.add_constant(df["x"])
m = sm.OLS(df["y"], X).fit()
print(f"OLS: coef={{m.params.iloc[1]:.6f}}, p={{m.pvalues.iloc[1]:.6f}}, R2={{m.rsquared:.4f}}")

dfe = df.copy()
for c in ["x","y"]: dfe[f"{{c}}_dm"] = dfe[c] - dfe.groupby("country")[c].transform("mean")
Xfe = sm.add_constant(dfe["x_dm"])
fem = sm.OLS(dfe["y_dm"], Xfe).fit()
print(f"FE:  coef={{fem.params.iloc[-1]:.6f}}, p={{fem.pvalues.iloc[-1]:.6f}}, R2w={{fem.rsquared:.4f}}")

pr,pp = stats.pearsonr(df["x"], df["y"])
sr,sp = stats.spearmanr(df["x"], df["y"])
print(f"Pearson: r={{pr:.4f}} (p={{pp:.6f}})")
print(f"Spearman: rho={{sr:.4f}} (p={{sp:.6f}})")
print("Done.")
'''
        else:
            # ── World Bank reproduction script (original) ──
            script = f'''#!/usr/bin/env python3
"""
Reproduction Script - Generated by Empirica
Hypothesis: {plan["statement"]}
Data source: World Bank World Development Indicators
Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}

Run: pip install pandas statsmodels scipy requests
     python reproduce.py
"""
import requests, pandas as pd, numpy as np
import statsmodels.api as sm, scipy.stats as stats

X_IND = "{plan.get("independent_var", "")}"
Y_IND = "{plan.get("dependent_var", "")}"
CONTROLS = [
{controls_code}]
START, END = {plan.get("start_year", 2000)}, {plan.get("end_year", 2023)}
WINSORIZE = {review.get("winsorize", False)}
WIN_PCT = {review.get("winsorize_percentile", 1)} / 100
MIN_OBS = {review.get("min_observations_per_country", 3)}

AGGREGATES = {{"WLD","HIC","LIC","LMC","MIC","UMC","LMY","HPC","EAS","ECS","LCN","MEA","NAC","SAS","SSF"}}

def fetch_wb(ind, s, e):
    rows, page = [], 1
    while True:
        r = requests.get(f"https://api.worldbank.org/v2/country/all/indicator/{{ind}}?date={{s}}:{{e}}&format=json&per_page=1000&page={{page}}", timeout=30).json()
        if len(r) < 2 or not r[1]: break
        for rec in r[1]:
            v = rec.get("value")
            if v is not None:
                cc = rec["country"]["id"]
                if cc not in AGGREGATES:
                    rows.append({{"country": rec["country"]["value"], "cc": cc, "year": int(rec["date"]), "value": float(v)}})
        if page >= r[0].get("pages", 1): break
        page += 1
    return pd.DataFrame(rows)

print("Fetching data...")
xd = fetch_wb(X_IND, START, END)
yd = fetch_wb(Y_IND, START, END)
df = xd.rename(columns={{"value":"x"}}).merge(yd.rename(columns={{"value":"y"}})[["country","year","y"]], on=["country","year"])

if WINSORIZE:
    for c in ["x","y"]:
        df[c] = df[c].clip(df[c].quantile(WIN_PCT), df[c].quantile(1-WIN_PCT))
counts = df.groupby("country").size()
df = df[df["country"].isin(counts[counts >= MIN_OBS].index)]
print(f"Data: {{len(df)}} obs, {{df['country'].nunique()}} countries")

X = sm.add_constant(df["x"])
m = sm.OLS(df["y"], X).fit()
print(f"OLS: coef={{m.params.iloc[1]:.6f}}, p={{m.pvalues.iloc[1]:.6f}}, R2={{m.rsquared:.4f}}")

dfe = df.copy()
for c in ["x","y"]: dfe[f"{{c}}_dm"] = dfe[c] - dfe.groupby("country")[c].transform("mean")
Xfe = sm.add_constant(dfe["x_dm"])
fem = sm.OLS(dfe["y_dm"], Xfe).fit()
print(f"FE:  coef={{fem.params.iloc[-1]:.6f}}, p={{fem.pvalues.iloc[-1]:.6f}}, R2w={{fem.rsquared:.4f}}")

pr,pp = stats.pearsonr(df["x"], df["y"])
sr,sp = stats.spearmanr(df["x"], df["y"])
print(f"Pearson: r={{pr:.4f}} (p={{pp:.6f}})")
print(f"Spearman: rho={{sr:.4f}} (p={{sp:.6f}})")
print("Done.")
'''
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        with open(output_path, "w") as f:
            f.write(script)
        print(f"  ✅ Script saved: {output_path}")


# ============================================================================
# MAIN PIPELINE
# ============================================================================
def run_empirica(hypothesis: str, output_dir: str = OUTPUT_DIR,
                 advocacy_angle: str = "", advocacy_temperature: int = 1):
    print("\n" + "=" * 60)
    print("  EMPIRICA v1.8.0")
    print("=" * 60)
    print(f"  Input: {hypothesis}")
    if advocacy_angle:
        print(f"  Advocacy: {advocacy_angle} (temperature: {advocacy_temperature}/10)")
    print("=" * 60)

    os.makedirs(output_dir, exist_ok=True)

    # Agent 1: Parse (extended thinking)
    plan = ai_parse_hypothesis(hypothesis)

    # Agent 2a: Fetch data (source-aware router)
    source = plan.get("data_source", "worldbank")
    print(f"\n📊 AGENT 2a: Fetching data (source: {source})...")

    if source == "ameco" and HAS_DBNOMICS:
        # ── AMECO via DBnomics ──
        dbn = DBnomicsFetcher()
        ameco_x = plan.get("ameco_independent", {})
        ameco_y = plan.get("ameco_dependent", {})
        if not ameco_x or not ameco_y:
            raise ValueError("AMECO selected but ameco_independent/dependent not provided by Agent 1.")
        x_data = dbn.fetch(ameco_x["dataset"], ameco_x.get("dimensions"), plan["start_year"], plan["end_year"])
        y_data = dbn.fetch(ameco_y["dataset"], ameco_y.get("dimensions"), plan["start_year"], plan["end_year"])
        if x_data.empty or y_data.empty:
            print("  ⚠️  AMECO fetch failed — falling back to World Bank...")
            source = "worldbank"  # fallback below
        else:
            # Use AMECO labels
            plan.setdefault("x_label", ameco_x.get("label", ameco_x["dataset"]))
            plan.setdefault("y_label", ameco_y.get("label", ameco_y["dataset"]))
            plan["_actual_source"] = "ameco"

    if source == "ameco" and not HAS_DBNOMICS:
        print("  ⚠️  dbnomics not installed — falling back to World Bank...")
        source = "worldbank"

    if source in ("worldbank", "both"):
        # ── World Bank (original path) ──
        wb = WorldBankFetcher()
        x_data = wb.fetch(plan["independent_var"], plan["start_year"], plan["end_year"])
        y_data = wb.fetch(plan["dependent_var"], plan["start_year"], plan["end_year"])
        if x_data.empty or y_data.empty:
            raise ValueError("Could not fetch data from World Bank.")
        plan["_actual_source"] = "worldbank"

    # ── SAME MERGE — downstream is source-agnostic ──
    df = x_data.rename(columns={"value": "x"}).merge(
        y_data.rename(columns={"value": "y"})[["country", "year", "y"]],
        on=["country", "year"],
    )

    # Fetch controls (from same source)
    controls_fetched = []
    for ctrl in plan.get("control_vars", []):
        ctrl_source = ctrl.get("source", plan.get("_actual_source", "worldbank"))
        if ctrl_source == "ameco" and HAS_DBNOMICS:
            dbn = DBnomicsFetcher()
            cd = dbn.fetch(ctrl["code"], ctrl.get("dimensions"), plan["start_year"], plan["end_year"])
        else:
            wb = WorldBankFetcher()
            cd = wb.fetch(ctrl["code"], plan["start_year"], plan["end_year"])
        if not cd.empty:
            cn = f"control_{ctrl['code'].replace('.', '_')}"
            cd = cd.rename(columns={"value": cn})
            df = df.merge(cd[["country", "year", cn]], on=["country", "year"], how="left")
            controls_fetched.append(ctrl)

    print(f"\n  ✅ Merged: {len(df)} rows, {df['country'].nunique()} countries")

    # Agent 2b: Literature (dual queries for 20+ papers)
    lit = LiteratureSearcher()
    literature = lit.search(plan)

    # Agent 3: Review data (extended thinking)
    review = ai_review_data(df, plan)
    df = apply_cleaning(df, review)
    if len(df) < 10:
        raise ValueError(f"Only {len(df)} observations after cleaning.")

    # Agent 4: Statistics
    engine = StatisticsEngine()
    results = engine.run_all(df, plan)
    if "error" in results:
        raise ValueError(f"Analysis failed: {results['error']}")

    # Generate charts
    scatterplot_path = ""
    coeff_plot_path = ""
    try:
        scatterplot_path = generate_scatterplot(df, plan, output_dir)
    except Exception as e:
        print(f"  ⚠️  Scatterplot failed: {e}")
    try:
        coeff_plot_path = generate_coefficient_plot(results, plan, output_dir)
    except Exception as e:
        print(f"  ⚠️  Coefficient plot failed: {e}")

    # Agent 5: Interpret (extended thinking)
    interpretation = ai_interpret_results(results, plan, advocacy_angle, advocacy_temperature)

    # Agent 6: Write (no extended thinking — just good prompts)
    writer = PaperWriter(plan, results, interpretation, literature, advocacy_angle, advocacy_temperature)
    sections = writer.write_all()

    # Agent 6b: Proofread (no extended thinking)
    sections = ai_proofread(sections)

    # Agent 7: Assemble
    paper_path = os.path.join(output_dir, "paper.docx")
    assembler = DocumentAssembler()
    assembler.create(plan, sections, results, literature, controls_fetched, paper_path,
                     scatterplot_path=scatterplot_path, coeff_plot_path=coeff_plot_path)

    repro_path = os.path.join(output_dir, "reproduce.py")
    ReproductionScriptGenerator().generate(plan, review, results, repro_path)

    print("\n" + "=" * 60)
    print("  ✅ EMPIRICA v1.8.0 COMPLETE")
    print("=" * 60)
    print(f"  Paper:  {paper_path}")
    print(f"  Code:   {repro_path}")
    main_r = results.get("ols_controls", results.get("ols", {}))
    print(f"  Result: B={main_r.get('coefficient','N/A')}, p={main_r.get('p_value','N/A')}, R2={main_r.get('r_squared','N/A')}")
    print("=" * 60)

    return results


# ============================================================================
# CLI ENTRY POINT
# ============================================================================
if __name__ == "__main__":
    if len(sys.argv) > 1:
        run_empirica(" ".join(sys.argv[1:]))
    else:
        print("Usage: python empirica_v3.py \"Your hypothesis here\"")
        sys.exit(1)
